#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
月报自动化处理脚本 v2
功能：根据模板中的业务系统自动筛选原始数据并填充
"""

import os
import sys
import re
from datetime import datetime
from collections import defaultdict

sys.stdout.reconfigure(encoding='utf-8')

# ============ 配置区域 ============
BASE_DIR = r'D:\月报自动化'
RAW_DATA_DIR = os.path.join(BASE_DIR, '月报原始数据')
TEMPLATE_DIR = os.path.join(BASE_DIR, '月报模板')
OUTPUT_DIR = os.path.join(BASE_DIR, '输出月报')

# 数据文件路径（在generate_report中根据目标月份动态设置）
RESOURCE_FILE = None
WORKORDER_FILE = None
FORTRESS_FILE = None
VPN_FILE = None

def setup_data_paths(year, month):
    """根据年月设置数据文件路径（自动检测文件格式）"""
    global RESOURCE_FILE, WORKORDER_FILE, FORTRESS_FILE, VPN_FILE
    month_str = f"{year}-{month:02d}"
    
    # 基础文件路径
    RESOURCE_FILE = os.path.join(RAW_DATA_DIR, f'{month_str}月报资源使用率详情列表.xls')
    
    # 工单支持多种命名和格式
    workorder_patterns = [
        f'{month_str}工单列表.xlsx',
        f'{month_str}工单总量.xlsx', 
        f'{month}月-工单列表.xlsx',
        f'{month}月-工单总量.xlsx',
    ]
    WORKORDER_FILE = None
    for pattern in workorder_patterns:
        path = os.path.join(RAW_DATA_DIR, pattern)
        if os.path.exists(path):
            WORKORDER_FILE = path
            break
    if WORKORDER_FILE is None:
        WORKORDER_FILE = os.path.join(RAW_DATA_DIR, workorder_patterns[0])
    
    FORTRESS_FILE = os.path.join(RAW_DATA_DIR, f'{month_str}-堡垒机.xlsx')
    
    # VPN支持三种格式：优先找 .csv，然后 .xlsx，最后 .xls
    vpn_patterns = [
        f'{month}月-VPN审计记录.csv',
        f'{month_str}vpn审计.xlsx',
        f'{month_str}vpn审计.xls',
    ]
    VPN_FILE = None
    for pattern in vpn_patterns:
        path = os.path.join(RAW_DATA_DIR, pattern)
        if os.path.exists(path):
            VPN_FILE = path
            break
    if VPN_FILE is None:
        VPN_FILE = os.path.join(RAW_DATA_DIR, vpn_patterns[1])

# ============ 快照备份功能 ============
def get_fridays_from_month(year, month):
    """获取指定年月的所有周五日期列表"""
    import calendar
    from datetime import date
    
    cal = calendar.monthcalendar(year, month)
    fridays = []
    for week in cal:
        friday_day = week[4]  # 4=周五（0=周一）
        if friday_day != 0:
            friday_date = date(year, month, friday_day)
            fridays.append(friday_date)
    return fridays

def get_last_month(year, month):
    """获取上个月的年月"""
    if month == 1:
        return year - 1, 12
    else:
        return year, month - 1

def validate_input_files(year, month):
    """验证所有输入文件是否存在"""
    setup_data_paths(year, month)
    
    required_files = {
        '资源数据': RESOURCE_FILE,
    }
    
    optional_files = {
        '工单数据': WORKORDER_FILE,
        '堡垒机数据': FORTRESS_FILE,
        'VPN数据': VPN_FILE,
    }
    
    missing_required = []
    for name, path in required_files.items():
        if not os.path.exists(path):
            missing_required.append(f"{name}: {path}")
    
    if missing_required:
        print(f"❌ 缺少必需文件:")
        for item in missing_required:
            print(f"   - {item}")
        return False
    
    # 检查可选文件
    existing_optional = []
    missing_optional = []
    for name, path in optional_files.items():
        if os.path.exists(path):
            existing_optional.append(name)
        else:
            missing_optional.append(name)
    
    print(f"✅ 必需文件检查通过")
    if existing_optional:
        print(f"📎 可选文件: {', '.join(existing_optional)}")
    if missing_optional:
        print(f"⚠️  可选文件缺失: {', '.join(missing_optional)}")
    
    return True

def generate_backup_records(hosts, backup_year, backup_month, backup_person="张昊"):
    """
    生成快照备份记录
    
    Args:
        hosts: 主机列表，每个主机包含 'host_name' 和 'ip'
        backup_year: 备份年份
        backup_month: 备份月份
        backup_person: 备份负责人（默认：张昊）
    
    Returns:
        备份记录列表，格式：[{seq, host_name, ip, backup_time, backup_type, person}, ...]
    """
    fridays = get_fridays_from_month(backup_year, backup_month)
    
    records = []
    
    # 按周五分组，每个周五为一波，每波序号从1开始
    for friday in fridays:
        # 备份时间格式: YYYY/M/D 22:00 (不补零)
        backup_time = f"{friday.year}/{friday.month}/{friday.day} 22:00"
        
        # 这一波的所有主机，序号从1开始
        for idx, host in enumerate(hosts):
            host_name = host.get('host_name', '')
            ip = host.get('ip', '')
            
            record = {
                'seq': idx + 1,  # 按周五分组，每波从1开始：1-121
                'host_name': host_name,
                'ip': ip,
                'backup_time': backup_time,
                'backup_type': '快照备份',  # 备份类型
                'person': backup_person  # 备份负责人
            }
            records.append(record)
    
    return records

# ============ Excel读取工具 ============
def read_resource_data(file_path, target_systems=None):
    """读取资源使用率数据，根据业务系统筛选
    
    返回:
        systems_data: 匹配的业务系统数据字典
        col_idx: 列索引信息
        matched_systems: 匹配到的业务系统名称集合
        unmatched_systems: 未匹配到的业务系统名称集合
    """
    try:
        import xlrd
        book = xlrd.open_workbook(file_path)
        sheet = book.sheet_by_index(0)
        
        print(f"\n📊 读取资源数据: {os.path.basename(file_path)}")
        print(f"   总行数: {sheet.nrows}")
        if target_systems:
            print(f"   目标业务系统: {target_systems}")
        
        # 用于跟踪匹配情况
        target_systems_set = set(target_systems) if target_systems else set()
        matched_systems = set()
        
        # 找到表头行
        header_row = None
        headers = []
        for i in range(min(15, sheet.nrows)):
            row_values = sheet.row_values(i)
            if any('业务系统' in str(cell) for cell in row_values):
                header_row = i
                headers = row_values
                print(f"   表头行: {i}, 列名: {headers[:10]}")
                break
        
        if not headers:
            print("   ⚠️ 未找到表头")
            return {}, {}
        
        # 找到关键列索引
        col_idx = {}
        for i, h in enumerate(headers):
            h_str = str(h).strip()
            if '业务系统名称' in h_str:
                col_idx['system'] = i
            elif '云主机名称' in h_str:
                col_idx['host_name'] = i
            elif '浮动IP' in h_str:
                col_idx['ip'] = i  # 使用浮动IP而不是固定IP
            elif h_str == 'CPU':
                col_idx['cpu'] = i
            elif '内存' in h_str and 'GB' in h_str:
                col_idx['memory'] = i
            elif '磁盘' in h_str and '总' in h_str:
                col_idx['storage'] = i
            elif 'CPU使用率' in h_str and 'AVG' in h_str:
                col_idx['cpu_usage'] = i
            elif 'CPU使用率' in h_str and 'MAX' in h_str:
                col_idx['cpu_usage_max'] = i
            elif '内存使用率' in h_str and 'AVG' in h_str:
                col_idx['mem_usage'] = i
            elif '内存使用率' in h_str and 'MAX' in h_str:
                col_idx['mem_usage_max'] = i
            elif '磁盘使用率' in h_str and 'AVG' in h_str:
                col_idx['disk_usage'] = i
        
        print(f"   列索引: {col_idx}")
        
        # 读取并筛选数据
        systems_data = defaultdict(lambda: {
            'hosts': [],
            'cpu_count': 0,
            'memory_gb': 0,
            'storage_gb': 0,
            'host_count': 0,
            'cpu_usage_values': [],
            'mem_usage_values': [],
            'disk_usage_values': []
        })
        
        for i in range(header_row + 1, sheet.nrows):
            row = sheet.row_values(i)
            if not row or len(row) < 3:
                continue
            
            system_name = str(row[col_idx.get('system', 1)]).strip() if col_idx.get('system', 1) < len(row) else ''
            
            # 只处理目标业务系统
            if not system_name:
                continue
            if target_systems and system_name not in target_systems:
                continue
            
            # 记录匹配到的业务系统
            if target_systems:
                matched_systems.add(system_name)
            
            # 提取主机信息
            def get_val(idx, default=0):
                try:
                    if idx and idx < len(row):
                        v = row[idx]
                        if isinstance(v, (int, float)):
                            return v
                        if str(v).strip():
                            return float(v)
                except:
                    pass
                return default
            
            host_info = {
                'host_name': str(row[col_idx.get('host_name', 3)]).strip() if col_idx.get('host_name', 3) < len(row) else '',
                'ip': str(row[col_idx.get('ip', 9)]).strip() if col_idx.get('ip', 9) < len(row) else '',  # 浮动IP在索引9
                'cpu': int(get_val(col_idx.get('cpu', 5))),
                'memory': int(get_val(col_idx.get('memory', 6))),
                'storage': int(get_val(col_idx.get('storage', 7))),
                'cpu_usage': get_val(col_idx.get('cpu_usage')),
                'cpu_usage_max': get_val(col_idx.get('cpu_usage_max')),
                'mem_usage': get_val(col_idx.get('mem_usage')),
                'mem_usage_max': get_val(col_idx.get('mem_usage_max')),
                'disk_usage': get_val(col_idx.get('disk_usage')),
            }
            
            systems_data[system_name]['hosts'].append(host_info)
            systems_data[system_name]['cpu_count'] += host_info['cpu']
            systems_data[system_name]['memory_gb'] += host_info['memory']
            systems_data[system_name]['storage_gb'] += host_info['storage']
            systems_data[system_name]['host_count'] += 1
            # 所有主机的使用率都参与计算（包括0值）
            systems_data[system_name]['cpu_usage_values'].append(host_info['cpu_usage'])
            systems_data[system_name]['mem_usage_values'].append(host_info['mem_usage'])
            systems_data[system_name]['disk_usage_values'].append(host_info['disk_usage'])
        
        # 计算平均使用率
        for sys_name, data in systems_data.items():
            data['cpu_usage_avg'] = sum(data['cpu_usage_values']) / len(data['cpu_usage_values']) if data['cpu_usage_values'] else 0
            data['mem_usage_avg'] = sum(data['mem_usage_values']) / len(data['mem_usage_values']) if data['mem_usage_values'] else 0
            data['disk_usage_avg'] = sum(data['disk_usage_values']) / len(data['disk_usage_values']) if data['disk_usage_values'] else 0
        
        # 计算未匹配的业务系统
        unmatched_systems = target_systems_set - matched_systems if target_systems else set()
        
        # 打印匹配统计
        if target_systems:
            print(f"\n   匹配统计:")
            print(f"   - 目标业务系统: {len(target_systems_set)} 个")
            print(f"   - 成功匹配: {len(matched_systems)} 个")
            print(f"   - 未匹配: {len(unmatched_systems)} 个")
            if unmatched_systems:
                print(f"   - 未匹配的系统: {sorted(unmatched_systems)}")
        
        return dict(systems_data), col_idx, matched_systems, unmatched_systems
        
    except Exception as e:
        print(f"❌ 读取资源数据失败: {e}")
        import traceback
        traceback.print_exc()
        return {}, {}, set(), set()

# ============ Word处理 ============
def extract_customer_name(docx_path, doc=None):
    """从Word模板中提取客户名称（委办局名称）"""
    if doc is None:
        from docx import Document
        doc = Document(docx_path)
    
    # 方法1: 从文件名中提取
    filename = os.path.basename(docx_path)
    # 匹配 "XXXX-YYYY年MM月-客户名称.docx" 的格式
    match = re.search(r'-(\d{4}年\d{1,2}月)-(.+?)\.docx$', filename)
    if match:
        customer = match.group(2).strip()
        if customer:
            print(f"   从文件名提取客户名称: {customer}")
            return customer
    
    # 方法2: 从段落中找到包含"共有"和"在"的段落
    for para in doc.paragraphs[:20]:  # 只看前20段
        text = para.text
        # 匹配类似 "XX年XX月，客户名称共有..."
        match = re.search(r'\d{4}年\d{1,2}月，([^，]+?)共有\d+个业务系统', text)
        if match:
            customer = match.group(1).strip()
            if customer:
                print(f"   从段落提取客户名称: {customer}")
                return customer
    
    # 方法3: 从标题样式中找
    for para in doc.paragraphs[:10]:
        if para.style.name.startswith('Heading') or para.style.name.startswith('标题'):
            text = para.text.strip()
            if text and '报告' in text:
                # 尝试提取报告标题前的名称
                parts = text.split('政务云服务')
                if len(parts) > 0 and parts[0]:
                    customer = parts[0].strip()
                    print(f"   从标题提取客户名称: {customer}")
                    return customer
    
    print("   ⚠️ 未能自动提取客户名称，使用默认值")
    return "北京市农林科学院"

def extract_target_systems(doc):
    """从Word模板中提取目标业务系统名称"""
    systems = set()
    
    # 方法1: 从表格中提取（业务系统资源使用情况统计表）
    for table in doc.tables[:5]:
        # 先检查表头，找到业务系统名称列
        header_row = None
        system_col_idx = None
        
        for row_idx, row in enumerate(table.rows[:2]):  # 只看前2行找表头
            cells = [cell.text.strip() for cell in row.cells]
            for col_idx, cell_text in enumerate(cells):
                if '业务系统名称' in cell_text or (cell_text == '业务系统' and col_idx > 0):
                    header_row = row_idx
                    system_col_idx = col_idx
                    break
            if system_col_idx is not None:
                break
        
        if system_col_idx is None:
            # 默认第二列
            system_col_idx = 1
        
        # 从数据行提取业务系统
        for row in table.rows[header_row+1 if header_row is not None else 1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) > system_col_idx:
                system_name = cells[system_col_idx]
                # 过滤条件
                if system_name and len(system_name) >= 2:
                    # 排除非业务系统的关键词
                    exclude_keywords = ['工作类型', '工单数量', '业务系统名称', '业务系统', '需求处理', '故障处理', '咨询', '其他', '总量', '使用率', '序号', '平均值', '主机数量', 'CPU总量', '内存总量', '存储总量']
                    if not any(kw in system_name for kw in exclude_keywords):
                        if not system_name.isdigit():
                            systems.add(system_name)
    
    # 方法2: 从段落中查找"其中"后面的业务系统名称
    for para in doc.paragraphs:
        text = para.text
        if '其中' in text and '业务系统' in text:
            match = re.search(r'其中\s+(.+?)(?=在|共|运行|，)', text)
            if match:
                names_part = match.group(1)
                for name in names_part.split():
                    name = name.strip()
                    if name and len(name) >= 2 and not name.isdigit():
                        # 排除关键词
                        if '需求' not in name and '处理' not in name:
                            systems.add(name)
    
    result = sorted(list(systems))
    print(f"   从模板中提取的业务系统: {result}")
    return result

def read_workorder_data(file_path, target_systems=None):
    """读取工单数据，按业务系统统计"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        
        # 读取"全部工单"工作表
        if '全部工单' not in wb.sheetnames:
            print(f"   ⚠️ 未找到'全部工单'工作表，可用工作表: {wb.sheetnames[:5]}")
            wb.close()
            return {}
        
        ws = wb['全部工单']
        print(f"   读取工作表: 全部工单")
        
        # 读取表头
        rows = list(ws.iter_rows(values_only=True, max_row=20))
        header_row = None
        headers = []
        for i, row in enumerate(rows):
            if any('业务系统' in str(cell) for cell in row if cell):
                header_row = i
                headers = row
                break
        
        if not headers:
            print("   ⚠️ 未找到表头")
            wb.close()
            return {}
        
        # 找到关键列索引
        col_idx = {}
        for i, h in enumerate(headers):
            if h is None:
                continue
            h_str = str(h).strip()
            if '业务系统' in h_str:
                col_idx['system'] = i
            elif '工单类型' in h_str or '类型' in h_str:
                col_idx['type'] = i
        
        print(f"   工单表列索引: {col_idx}")
        
        # 统计工单数据
        workorder_stats = defaultdict(lambda: defaultdict(int))
        total_count = 0
        
        for row in ws.iter_rows(values_only=True, min_row=header_row+2):
            if not row or len(row) < 3:
                continue
            
            system_name = str(row[col_idx.get('system', 6)]).strip() if col_idx.get('system', 6) < len(row) and row[col_idx.get('system', 6)] else ''
            workorder_type = str(row[col_idx.get('type', 2)]).strip() if col_idx.get('type', 2) < len(row) and row[col_idx.get('type', 2)] else '其他'
            
            if not system_name:
                continue
            
            # 只统计目标业务系统
            if target_systems and system_name not in target_systems:
                continue
            
            workorder_stats[system_name][workorder_type] += 1
            total_count += 1
        
        wb.close()
        
        print(f"   统计到工单总数: {total_count}")
        if workorder_stats:
            for sys_name, types in workorder_stats.items():
                print(f"   - {sys_name}: {dict(types)}")
        
        return dict(workorder_stats)
        
    except Exception as e:
        print(f"❌ 读取工单数据失败: {e}")
        import traceback
        traceback.print_exc()
        return {}

def clear_and_set_cell(cell, value):
    """彻底清除单元格内容并设置新值，保持居中+框线格式"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 1. 清除并设置新值
    cell.text = str(value)
    
    # 2. 居中对齐
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. 确保单元格有框线
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    
    # 获取或创建 tcPr (table cell properties)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 创建表格边框
    try:
        tcBorders = parse_xml(r'''
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tcBorders>
        ''')
        tcPr.append(tcBorders)
    except Exception:
        pass  # 如果失败，不影响主要内容

def format_table_row(row):
    """格式化整行表格，确保所有单元格有居中和框线"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    
    for cell in row.cells:
        # 居中对齐
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加框线
        try:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                </w:tcBorders>
            ''')
            tcPr.append(tcBorders)
        except Exception:
            pass

def generate_report(template_path, output_path, target_year=None, target_month=None):
    """生成月报"""
    # 默认使用上个月
    if target_year is None or target_month is None:
        target_year, target_month = get_default_target_month()
    
    print(f"\n{'='*60}")
    print(f"🚀 月报自动化处理系统 v3.0")
    print(f"{'='*60}")
    print(f"目标月份: {target_year}年{target_month}月")
    print(f"处理模板: {os.path.basename(template_path)}")
    
    # 验证输入文件
    if not validate_input_files(target_year, target_month):
        return False
    
    print(f"{'='*60}")
    
    from docx import Document
    doc = Document(template_path)
    
    # 初始化变量
    original_customer = None
    
    # 0. 提取客户名称（委办局名称）
    print("\n步骤0: 提取客户名称...")
    customer_name = extract_customer_name(template_path, doc)
    print(f"   客户名称: {customer_name}")
    
    # 查找原文中的客户名称（用于替换）
    for para in doc.paragraphs[:20]:
        match = re.search(r'(\d{4}年\d{1,2}月)，([^，]+?)共有\d+个业务系统', para.text)
        if match:
            original_customer = match.group(2).strip()
            print(f"   原文客户名称: {original_customer}")
            break
    
    # 1. 提取目标业务系统
    print("\n步骤1: 提取模板中的业务系统...")
    target_systems = extract_target_systems(doc)
    if not target_systems:
        print("❌ 未能从模板中提取业务系统")
        return False
    
    print(f"   目标业务系统: {target_systems}")
    
    # 用于存储所有匹配状态信息（后续同步给用户）
    system_match_status = {
        'target_systems': list(target_systems),
        'resource_matched': [],
        'resource_unmatched': [],
        'workorder_matched': [],
        'workorder_unmatched': [],
    }
    
    # 2. 读取资源数据（仅筛选目标业务系统）
    print("\n步骤2: 读取资源数据（仅目标业务系统）...")
    resource_data, col_idx, resource_matched, resource_unmatched = read_resource_data(RESOURCE_FILE, target_systems)
    
    # 记录资源数据匹配状态
    system_match_status['resource_matched'] = list(resource_matched)
    system_match_status['resource_unmatched'] = list(resource_unmatched)
    
    if not resource_data:
        print("❌ 未找到匹配的资源数据")
        return False
    
    # 打印匹配结果
    print(f"\n   匹配到的业务系统数据:")
    total_hosts, total_cpu, total_mem, total_storage = 0, 0, 0, 0
    for sys_name, data in resource_data.items():
        print(f"   - {sys_name}: {data['host_count']}台主机, {data['cpu_count']}核CPU, {data['memory_gb']}GB内存, {data['storage_gb']}GB存储")
        total_hosts += data['host_count']
        total_cpu += data['cpu_count']
        total_mem += data['memory_gb']
        total_storage += data['storage_gb']
    
    # 2.5. 读取工单数据
    # 读取工单数据
    workorder_data = read_workorder_data(WORKORDER_FILE, target_systems)
    total_workorders = sum(sum(types.values()) for types in workorder_data.values()) if workorder_data else 0
    
    # 简洁的数据总结
    print(f"\n   数据汇总: {len(target_systems)}个业务系统 | {total_hosts}台主机 | {total_workorders}个工单")
    
    # 3. 替换日期和客户名称
    print("\n步骤3: 替换日期和客户名称...")
    # 获取上个月信息用于报告中的对比日期
    last_year, last_month = get_last_month(target_year, target_month)
    last_month_str = f"{last_year}年{last_month}月"
    target_month_str = f"{target_year}年{target_month}月"
    
    # 查找并保存原文中的客户名称（用于替换）
    original_customer = None
    for para in doc.paragraphs[:30]:
        match = re.search(r'(\d{4}年\d{1,2}月)，([^，]+?)共有\d+个业务系统', para.text)
        if match:
            original_customer = match.group(2).strip()
            break
    
    if original_customer and original_customer != customer_name:
        print(f"   检测到需要替换客户名称: {original_customer} -> {customer_name}")
    
    # 替换所有段落中的日期（包括目录页）
    replace_count = 0
    for para in doc.paragraphs:
        full_text = ''.join([run.text for run in para.runs])
        new_text = full_text
        
        # 替换各种日期格式为目标月份
        # 匹配 YYYY年M月 或 YYYY年MM月 格式
        date_patterns = [
            r'\d{4}年\d{1,2}月',
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, new_text)
            for match in matches:
                # 替换为目标月份
                new_text = new_text.replace(match, target_month_str)
        
        # 如果文本有变化，更新run
        if new_text != full_text:
            # 清除所有run，只在第一个run设置新文本
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = new_text
                else:
                    run.text = ''
            replace_count += 1
        
        # 替换客户名称（如果不同）
        if original_customer and original_customer != customer_name and original_customer in full_text:
            for run in para.runs:
                run.text = run.text.replace(original_customer, customer_name)
    
    print(f"   共替换 {replace_count} 个段落中的日期")
    
    # 4. 替换基础资源台账
    print("\n步骤4: 更新基础资源台账概况...")
    for para in doc.paragraphs:
        text = para.text
        if '共有' in text and '个业务系统' in text:
            # 收集所有run的文本
            full_text = ''.join([run.text for run in para.runs])
            
            # 执行替换
            new_text = re.sub(r'共有\d+个业务系统', f'共有{len(target_systems)}个业务系统', full_text)
            new_text = re.sub(r'共使用\d+台主机', f'共使用{total_hosts}台主机', new_text)
            new_text = re.sub(r'\d+颗CPU', f'{total_cpu}颗CPU', new_text)
            new_text = re.sub(r'\d+GB内存', f'{total_mem}GB内存', new_text)
            new_text = re.sub(r'\d+GB存储', f'{total_storage}GB存储', new_text)
            
            # 如果客户名称不同，也替换
            if original_customer and original_customer != customer_name:
                new_text = new_text.replace(original_customer, customer_name)
            
            if new_text != full_text:
                # 清除所有run，只在第一个run设置新文本
                for i, run in enumerate(para.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ''
                print(f"   更新: {new_text[:60]}...")
    
    # 4.5. 更新"本月技术支撑统计"表格
    print("\n步骤4.5: 更新本月技术支撑统计表格...")
    for i, table in enumerate(doc.tables[:3]):
        # 检查是否是技术支撑统计表（表1）
        # 表格0通常是：['序号', '工作类型', '工单数量']
        if len(table.rows) >= 2 and len(table.columns) == 3:
            first_row = [cell.text.strip() for cell in table.rows[0].cells]
            if first_row == ['序号', '工作类型', '工单数量'] or '工作类型' in first_row:
                print(f"   找到技术支撑统计表 (表格{i})")
                
                # 获取最后一行（数据行）
                data_row = table.rows[1]
                cells = data_row.cells
                
                if len(cells) >= 3:
                    # 使用clear_and_set_cell更新工单数量
                    clear_and_set_cell(cells[2], total_workorders)
                    print(f"   更新工单数量: {total_workorders}")
                    print(f"   已设置居中对齐")
                    break  # 只处理第一个匹配的表格
    
    # 4.6. 更新"业务系统资源使用情况统计"表格（CPU、内存、存储）
    print("\n步骤4.6: 更新业务系统资源使用情况统计表格...")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 遍历表格查找CPU、内存、存储使用率表
    for i, table in enumerate(doc.tables):
        if len(table.rows) < 2 or len(table.columns) < 5:
            continue
        
        # 检查表头
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        first_header = header_cells[0] if header_cells else ''
        
        # 确定表格类型
        table_type = None
        if 'CPU使用率' in first_header or ('CPU' in first_header and '使用率' in first_header):
            table_type = 'CPU'
        elif '内存使用率' in first_header or ('内存' in first_header and '使用率' in first_header):
            table_type = '内存'
        elif '磁盘使用率' in first_header or ('磁盘' in first_header and '使用率' in first_header) or ('存储' in first_header and '使用率' in first_header):
            table_type = '存储'
        
        if not table_type:
            continue
        
        print(f"   找到 {table_type} 使用率表 (表格{i})")
        
        # 找到真正包含列名的表头行（包含"业务系统"和"序号"的行）
        header_row_idx = 0
        for r_idx, r in enumerate(table.rows[:3]):
            r_text = ' '.join([c.text.strip() for c in r.cells])
            if '业务系统' in r_text and ('序号' in r_text or '主机' in r_text):
                header_row_idx = r_idx
                break
        
        header_row = table.rows[header_row_idx]
        headers = [cell.text.strip() for cell in header_row.cells]
        # 找列索引
        col_map = {}
        for idx, h in enumerate(headers):
            if '序号' in h or h.isdigit():
                col_map['seq'] = idx
            elif '业务系统' in h:
                col_map['system'] = idx
            elif '主机' in h and '数量' in h:
                col_map['host_count'] = idx
            elif table_type == 'CPU' and (h == 'CPU' or 'CPU总量' in h or 'CPU总数' in h):
                col_map['total'] = idx
            elif table_type == '内存' and ('内存' in h and ('总量' in h or '总数' in h)):
                col_map['total'] = idx
            elif table_type == '存储' and (('磁盘' in h and ('总量' in h or '总数' in h)) or ('存储' in h and ('总量' in h or '总数' in h))):
                col_map['total'] = idx
            elif '使用率' in h:
                col_map['usage'] = idx
        
        # 找到数据起始行
        data_start_row = header_row_idx + 1
        
        for row_idx in range(data_start_row, len(table.rows)):
            row = table.rows[row_idx]
            cells = row.cells
            
            if len(cells) < 4:
                continue
            
            # 获取业务系统名称
            system_name = cells[col_map.get('system', 1)].text.strip() if col_map.get('system', 1) < len(cells) else ''
            
            # 跳过表头行或空行
            if not system_name or system_name in ['业务系统名称', '业务系统', '序号']:
                continue
            
            # 在resource_data中查找该业务系统的数据
            if system_name in resource_data:
                data = resource_data[system_name]
                
                # 更新主机数量
                if 'host_count' in col_map:
                    clear_and_set_cell(cells[col_map['host_count']], data['host_count'])
                
                # 更新总量
                if 'total' in col_map:
                    # 更新总量
                    if table_type == 'CPU':
                        clear_and_set_cell(cells[col_map['total']], data['cpu_count'])
                    elif table_type == '内存':
                        clear_and_set_cell(cells[col_map['total']], data['memory_gb'])
                    elif table_type == '存储':
                        clear_and_set_cell(cells[col_map['total']], data['storage_gb'])
                
                # 更新平均使用率
                if 'usage' in col_map:
                    if table_type == 'CPU':
                        usage = data.get('cpu_usage_avg', 0)
                    elif table_type == '内存':
                        usage = data.get('mem_usage_avg', 0)
                    elif table_type == '存储':
                        usage = data.get('disk_usage_avg', 0)
                    
                    clear_and_set_cell(cells[col_map['usage']], f"{usage:.2f}")
                
                print(f"      更新 {system_name}: {data['host_count']}台, 使用率{usage:.2f}%")
    
    print(f"   资源使用情况表格更新完成")
    
    # 更新各附件表格
    print("\n步骤5: 更新附件1 - 政务云基础资源台账...")
    update_attachment1(doc, resource_data, total_hosts)
    
    print("\n步骤6: 更新附件2 - 资源使用率详情...")
    update_attachment2(doc, resource_data)
    
    print("\n步骤7: 更新附件3 - 安全审计记录...")
    update_attachment3(doc, resource_data)
    
    print("\n步骤8: 更新快照备份服务报告...")
    update_snapshot_backup(doc, resource_data, target_year, target_month)
    
    print("\n步骤9: 更新网页防篡改服务报告...")
    update_web_tamper_report(doc, target_year, target_month)
    
    # 10. 保存输出
    print("\n步骤10: 保存生成结果...")
    try:
        doc.save(output_path)
        print(f"   ✅ 已保存: {output_path}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        return False
    
    # 11. 同步匹配结果
    print("\n" + "="*60)
    print("📋 业务系统匹配情况")
    print("="*60)
    
    # 6.1 资源数据匹配情况
    print(f"\n📊 资源数据匹配情况:")
    print(f"   目标业务系统总数: {len(system_match_status['target_systems'])} 个")
    print(f"   ✅ 成功匹配: {len(system_match_status['resource_matched'])} 个")
    if system_match_status['resource_matched']:
        for sys in sorted(system_match_status['resource_matched']):
            print(f"      - {sys}")
    
    print(f"   ❌ 未匹配: {len(system_match_status['resource_unmatched'])} 个")
    if system_match_status['resource_unmatched']:
        for sys in sorted(system_match_status['resource_unmatched']):
            print(f"      - {sys}")
    
    if system_match_status['resource_unmatched']:
        print("\n   ⚠️ 提示：未匹配的业务系统可能是名称不一致导致")
    
    print("\n" + "="*60)
    
    return True

def update_attachment1(doc, resource_data, total_hosts):
    """更新附件1 - 政务云基础资源台账表格"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    print(f"   主机总数: {total_hosts}")
    
    # 找到附件1表格（通常是主机数量最多的那个详细表格）
    target_table = None
    target_table_idx = None
    
    for i, table in enumerate(doc.tables):
        # 跳过已处理的表格（前几个是统计表）
        if i < 4:
            continue
        
        # 检查是否是附件1（有7列：序号、业务系统、主机名、IP、CPU、内存、存储）
        if len(table.columns) >= 6:
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            header_str = ' '.join(header_cells)
            
            # 判断是否是附件1表格
            if ('序号' in header_str and '业务系统' in header_str and 
                ('主机IP' in header_str or 'IP' in header_str) and
                ('CPU' in header_str or '内存' in header_str)):
                
                # 检查行数是否匹配（或接近）
                if len(table.rows) > 10 or len(table.rows) >= total_hosts:
                    target_table = table
                    target_table_idx = i
                    print(f"   找到附件1表格: {len(table.columns)}列 x {len(table.rows)}行")
                    break
    
    if not target_table:
        print("   ⚠️ 未找到附件1表格")
        return
    
    # 分析表头，确定列索引
    headers = [cell.text.strip() for cell in target_table.rows[0].cells]
    col_map = {}
    
    for idx, h in enumerate(headers):
        h_lower = h.lower()
        if '序号' in h or h == '序号':
            col_map['seq'] = idx
        elif '业务系统' in h:
            col_map['system'] = idx
        elif '主机名称' in h or '云主机' in h or '主机名' in h:
            col_map['host_name'] = idx
        elif 'IP' in h or 'ip' in h:
            col_map['ip'] = idx
        elif h == 'CPU' or 'cpu' in h_lower:
            col_map['cpu'] = idx
        elif '内存' in h:
            col_map['memory'] = idx
        elif '磁盘' in h or '存储' in h:
            col_map['storage'] = idx
    
    # 收集所有需要填充的主机数据
    all_hosts = []
    seq = 1
    for system_name in sorted(resource_data.keys()):
        data = resource_data[system_name]
        for host in data['hosts']:
            all_hosts.append({
                'seq': seq,
                'system': system_name,
                'host_name': host['host_name'],
                'ip': host['ip'],
                'cpu': host['cpu'],
                'memory': host['memory'],
                'storage': host['storage']
            })
            seq += 1
    
    # 更新或添加行
    existing_rows = len(target_table.rows) - 1  # 减去表头行
    
    # 首先更新现有行
    rows_to_update = min(len(all_hosts), existing_rows)
    for i in range(rows_to_update):
        row = target_table.rows[i + 1]  # +1 跳过表头
        host = all_hosts[i]
        cells = row.cells
        
        # 更新各列（使用clear_and_set_cell确保居中+框线）
        if 'seq' in col_map and col_map['seq'] < len(cells):
            clear_and_set_cell(cells[col_map['seq']], str(host['seq']))
        if 'system' in col_map and col_map['system'] < len(cells):
            clear_and_set_cell(cells[col_map['system']], host['system'])
        if 'host_name' in col_map and col_map['host_name'] < len(cells):
            clear_and_set_cell(cells[col_map['host_name']], host['host_name'])
        if 'ip' in col_map and col_map['ip'] < len(cells):
            clear_and_set_cell(cells[col_map['ip']], host['ip'])
        if 'cpu' in col_map and col_map['cpu'] < len(cells):
            clear_and_set_cell(cells[col_map['cpu']], str(int(host['cpu'])))
        if 'memory' in col_map and col_map['memory'] < len(cells):
            clear_and_set_cell(cells[col_map['memory']], str(int(host['memory'])))
        if 'storage' in col_map and col_map['storage'] < len(cells):
            clear_and_set_cell(cells[col_map['storage']], str(int(host['storage'])))
    
    # 添加新行（如需要）
    if len(all_hosts) > existing_rows:
        for i in range(existing_rows, len(all_hosts)):
            host = all_hosts[i]
            # 添加新行并格式化
            new_row = target_table.add_row()
            format_table_row(new_row)  # 确保格式一致（居中+框线）
            cells = new_row.cells
            
            # 填充数据（全部使用clear_and_set_cell确保居中+框线）
            if 'seq' in col_map and col_map['seq'] < len(cells):
                clear_and_set_cell(cells[col_map['seq']], str(host['seq']))
            if 'system' in col_map and col_map['system'] < len(cells):
                clear_and_set_cell(cells[col_map['system']], host['system'])
            if 'host_name' in col_map and col_map['host_name'] < len(cells):
                clear_and_set_cell(cells[col_map['host_name']], host['host_name'])
            if 'ip' in col_map and col_map['ip'] < len(cells):
                clear_and_set_cell(cells[col_map['ip']], host['ip'])
            if 'cpu' in col_map and col_map['cpu'] < len(cells):
                clear_and_set_cell(cells[col_map['cpu']], str(int(host['cpu'])))
            if 'memory' in col_map and col_map['memory'] < len(cells):
                clear_and_set_cell(cells[col_map['memory']], str(int(host['memory'])))
            if 'storage' in col_map and col_map['storage'] < len(cells):
                clear_and_set_cell(cells[col_map['storage']], str(int(host['storage'])))
    
    # 删除多余行（如需要）
    elif len(all_hosts) < existing_rows:
        rows_to_delete = existing_rows - len(all_hosts)
        rows_to_delete = existing_rows - len(all_hosts)
        for _ in range(rows_to_delete):
            if len(target_table.rows) > 2:  # 保留表头和至少一行
                # 获取最后一行并删除
                row_to_delete = target_table.rows[-1]
                # 使用XML删除行
                tbl = target_table._tbl
                tr = row_to_delete._tr
                tbl.remove(tr)
    
    print(f"   附件1更新完成: 共 {len(all_hosts)} 台主机")

def update_attachment2(doc, resource_data):
    """更新附件2 - 资源使用率详情表格（CPU、内存、磁盘三个表）"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 遍历表格查找附件2的三个表
    attachment2_tables = []
    
    for i, table in enumerate(doc.tables):
        # 跳过前几个已处理的表格
        if len(table.rows) < 2 or len(table.columns) < 6:
            continue
        
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        header_str = ' '.join(headers)
        
        # 判断表格类型
        table_type = None
        if 'CPU' in header_str and '使用率' in header_str and ('平均' in header_str or '最大' in header_str):
            table_type = 'CPU'
        elif '内存' in header_str and '使用率' in header_str and ('平均' in header_str or '最大' in header_str):
            table_type = '内存'
        elif ('磁盘' in header_str or '存储' in header_str) and '使用率' in header_str:
            table_type = '磁盘'
        
        if table_type:
            attachment2_tables.append((i, table, table_type))
    
    print(f"   找到 {len(attachment2_tables)} 个使用率详情表")
    
    for table_idx, table, table_type in attachment2_tables:
        print(f"   更新 {table_type} 使用率详情表...")
        
        col_map = {}
        for idx, h in enumerate(headers):
            h_lower = h.lower()
            if '序号' in h:
                col_map['seq'] = idx
            elif '业务系统' in h:
                col_map['system'] = idx
            elif '主机名称' in h or '云主机' in h:
                col_map['host_name'] = idx
            elif 'IP' in h or 'ip' in h_lower:
                col_map['ip'] = idx
            elif table_type == 'CPU' and ('cpu' in h_lower or 'CPU' in h):
                col_map['resource'] = idx
            elif table_type == '内存' and '内存' in h:
                col_map['resource'] = idx
            elif table_type == '磁盘' and ('磁盘' in h or '存储' in h):
                col_map['resource'] = idx
            elif '平均' in h and '使用率' in h:
                col_map['avg_usage'] = idx
            elif '最大' in h and '使用率' in h:
                col_map['max_usage'] = idx
            elif '当前' in h and '使用率' in h:
                col_map['cur_usage'] = idx
        

        
        # 收集该类型所有主机的详细数据
        all_host_details = []
        seq = 1
        
        for system_name in sorted(resource_data.keys()):
            data = resource_data[system_name]
            for host in data['hosts']:
                host_detail = {
                    'seq': seq,
                    'system': system_name,
                    'host_name': host['host_name'],
                    'ip': host['ip'],
                    'resource': None,
                    'avg_usage': 0,
                    'max_usage': 0,
                    'cur_usage': 0
                }
                
                # 根据表类型设置资源值和使用率
                if table_type == 'CPU':
                    host_detail['resource'] = host['cpu']
                    host_detail['avg_usage'] = host.get('cpu_usage', 0)
                    # 如果有最大使用率数据可以使用，否则用平均值
                    host_detail['max_usage'] = host.get('cpu_usage_max', host.get('cpu_usage', 0))
                elif table_type == '内存':
                    host_detail['resource'] = host['memory']
                    host_detail['avg_usage'] = host.get('mem_usage', 0)
                    host_detail['max_usage'] = host.get('mem_usage_max', host.get('mem_usage', 0))
                elif table_type == '磁盘':
                    host_detail['resource'] = host['storage']
                    host_detail['cur_usage'] = host.get('disk_usage', 0)
                
                all_host_details.append(host_detail)
                seq += 1
        
        # 更新或添加行
        existing_rows = len(table.rows) - 1
        
        # 更新现有行
        rows_to_update = min(len(all_host_details), existing_rows)
        for i in range(rows_to_update):
            row = table.rows[i + 1]
            host = all_host_details[i]
            cells = row.cells
            
            if 'seq' in col_map and col_map['seq'] < len(cells):
                clear_and_set_cell(cells[col_map['seq']], str(host['seq']))
            if 'system' in col_map and col_map['system'] < len(cells):
                clear_and_set_cell(cells[col_map['system']], host['system'])
            if 'host_name' in col_map and col_map['host_name'] < len(cells):
                clear_and_set_cell(cells[col_map['host_name']], host['host_name'])
            if 'ip' in col_map and col_map['ip'] < len(cells):
                clear_and_set_cell(cells[col_map['ip']], host['ip'])
            if 'resource' in col_map and col_map['resource'] < len(cells):
                clear_and_set_cell(cells[col_map['resource']], str(int(host['resource'])))
            
            # 更新使用率
            if 'avg_usage' in col_map and col_map['avg_usage'] < len(cells):
                clear_and_set_cell(cells[col_map['avg_usage']], f"{host['avg_usage']:.2f}")
            if 'max_usage' in col_map and col_map['max_usage'] < len(cells):
                clear_and_set_cell(cells[col_map['max_usage']], f"{host['max_usage']:.2f}")
            if 'cur_usage' in col_map and col_map['cur_usage'] < len(cells):
                clear_and_set_cell(cells[col_map['cur_usage']], f"{host['cur_usage']:.2f}")
        
        # 添加新行
        if len(all_host_details) > existing_rows:
            for i in range(existing_rows, len(all_host_details)):
                host = all_host_details[i]
                new_row = table.add_row()
                format_table_row(new_row)  # 确保格式一致
                cells = new_row.cells
                
                if 'seq' in col_map and col_map['seq'] < len(cells):
                    clear_and_set_cell(cells[col_map['seq']], str(host['seq']))
                if 'system' in col_map and col_map['system'] < len(cells):
                    clear_and_set_cell(cells[col_map['system']], host['system'])
                if 'host_name' in col_map and col_map['host_name'] < len(cells):
                    clear_and_set_cell(cells[col_map['host_name']], host['host_name'])
                if 'ip' in col_map and col_map['ip'] < len(cells):
                    clear_and_set_cell(cells[col_map['ip']], host['ip'])
                if 'resource' in col_map and col_map['resource'] < len(cells):
                    clear_and_set_cell(cells[col_map['resource']], str(int(host['resource'])))
                if 'avg_usage' in col_map and col_map['avg_usage'] < len(cells):
                    clear_and_set_cell(cells[col_map['avg_usage']], f"{host['avg_usage']:.2f}")
                if 'max_usage' in col_map and col_map['max_usage'] < len(cells):
                    clear_and_set_cell(cells[col_map['max_usage']], f"{host['max_usage']:.2f}")
                if 'cur_usage' in col_map and col_map['cur_usage'] < len(cells):
                    clear_and_set_cell(cells[col_map['cur_usage']], f"{host['cur_usage']:.2f}")
        
        # 删除多余行
        elif len(all_host_details) < existing_rows:
            rows_to_delete = existing_rows - len(all_host_details)
            for _ in range(rows_to_delete):
                if len(table.rows) > 2:
                    tbl = table._tbl
                    tr = table.rows[-1]._tr
                    tbl.remove(tr)
        
        print(f"      {table_type}使用率详情表更新完成: 共 {len(all_host_details)} 台主机")
    
    print(f"\n   附件2更新完成")

def read_fortress_data(file_path, target_ips):
    """读取堡垒机审计数据，根据IP筛选"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        
        # 使用第一个工作表
        ws = wb[wb.sheetnames[0]]
        
        # 读取表头
        rows = list(ws.iter_rows(values_only=True, max_row=10))
        header_row = None
        headers = []
        for i, row in enumerate(rows):
            if any('资产IP' in str(cell) for cell in row if cell):
                header_row = i
                headers = row
                break
        
        if not headers:
            print("   ⚠️ 未找到堡垒机表头")
            wb.close()
            return []
        
        # 找到列索引
        col_map = {}
        for idx, h in enumerate(headers):
            if h is None:
                continue
            h_str = str(h).strip()
            if '用户账号' in h_str or '登录名' in h_str:
                col_map['user'] = idx
            elif '开始时间' in h_str:
                col_map['start_time'] = idx
            elif '结束时间' in h_str:
                col_map['end_time'] = idx
            elif '会话时长' in h_str:
                col_map['duration'] = idx
            elif '来自' in h_str or '来源IP' in h_str:
                col_map['source'] = idx
            elif '资产名' in h_str or '主机名' in h_str:
                col_map['asset_name'] = idx
            elif '资产IP' in h_str or '设备IP' in h_str:
                col_map['asset_ip'] = idx
            elif '协议' in h_str:
                col_map['protocol'] = idx
            elif '客户端类型' in h_str or '客户端' in h_str:
                col_map['client_type'] = idx
        
        # 读取并筛选数据
        fortress_records = []
        for row in ws.iter_rows(values_only=True, min_row=header_row+2):
            if not row or len(row) < 5:
                continue
            
            # 获取资产IP
            asset_ip = str(row[col_map.get('asset_ip', 5)]).strip() if col_map.get('asset_ip', 5) < len(row) and row[col_map.get('asset_ip', 5)] else ''
            
            if not asset_ip:
                continue
            
            # 检查IP是否在目标列表中
            if asset_ip not in target_ips:
                continue
            
            # 提取记录并处理时间格式（去掉微秒）
            def format_time(val):
                """格式化时间，去掉微秒"""
                if not val:
                    return ''
                time_str = str(val).strip()
                # 去掉微秒部分（.后面的数字）
                if '.' in time_str:
                    time_str = time_str.split('.')[0]
                return time_str
            
            record = {
                'start_time': format_time(row[col_map.get('start_time', 0)]) if col_map.get('start_time', 0) < len(row) else '',
                'end_time': format_time(row[col_map.get('end_time', 1)]) if col_map.get('end_time', 1) < len(row) else '',
                'duration': str(row[col_map.get('duration', 2)]).strip() if col_map.get('duration', 2) < len(row) and row[col_map.get('duration', 2)] else '',
                'source': str(row[col_map.get('source', 3)]).strip() if col_map.get('source', 3) < len(row) and row[col_map.get('source', 3)] else '',
                'user': str(row[col_map.get('user', 4)]).strip() if col_map.get('user', 4) < len(row) and row[col_map.get('user', 4)] else '',
                'asset_name': str(row[col_map.get('asset_name', 5)]).strip() if col_map.get('asset_name', 5) < len(row) and row[col_map.get('asset_name', 5)] else '',
                'asset_ip': asset_ip,
                'protocol': str(row[col_map.get('protocol', 6)]).strip() if col_map.get('protocol', 6) < len(row) and row[col_map.get('protocol', 6)] else '',
                'client_type': str(row[col_map.get('client_type', 7)]).strip() if col_map.get('client_type', 7) < len(row) and row[col_map.get('client_type', 7)] else '',
            }
            
            fortress_records.append(record)
        
        wb.close()
        
        print(f"   匹配到 {len(fortress_records)} 条堡垒机审计记录")
        return fortress_records
        
    except Exception as e:
        print(f"❌ 读取堡垒机数据失败: {e}")
        import traceback
        traceback.print_exc()
        return []

def update_attachment3(doc, resource_data):
    """更新附件3 - 安全服务报告（堡垒机审计记录）"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 1. 提取所有主机IP
    all_host_ips = set()
    for system_name, data in resource_data.items():
        for host in data['hosts']:
            if host['ip']:
                all_host_ips.add(host['ip'])
    
    print(f"   从资源数据中提取到 {len(all_host_ips)} 个主机IP")
    
    if not all_host_ips:
        print("   ⚠️ 未找到主机IP，跳过堡垒机审计更新")
        return
    
    # 2. 查找堡垒机审计记录表格
    target_table = None
    for i, table in enumerate(doc.tables):
        # 检查表头
        if len(table.rows) < 2 or len(table.columns) < 8:
            continue
        
        header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if '开始时间' in header_text and '资产IP' in header_text and '协议' in header_text:
            target_table = table
            break
    
    if not target_table:
        print("   ⚠️ 未找到堡垒机审计记录表格")
        return
    
    # 3. 读取堡垒机数据
    fortress_records = read_fortress_data(FORTRESS_FILE, all_host_ips)
    
    # 4. 分析表头
    headers = [cell.text.strip() for cell in target_table.rows[0].cells]
    col_map = {}
    for idx, h in enumerate(headers):
        if '开始时间' in h:
            col_map['start_time'] = idx
        elif '结束时间' in h:
            col_map['end_time'] = idx
        elif '会话时长' in h:
            col_map['duration'] = idx
        elif '来自' in h or '来源' in h:
            col_map['source'] = idx
        elif '资产名' in h:
            col_map['asset_name'] = idx
        elif '资产IP' in h:
            col_map['asset_ip'] = idx
        elif '协议' in h:
            col_map['protocol'] = idx
        elif '客户端类型' in h or '客户端' in h:
            col_map['client_type'] = idx
    
    # 5. 更新表格数据
    existing_rows = len(target_table.rows) - 1  # 减去表头
    
    # 更新现有行
    rows_to_update = min(len(fortress_records), existing_rows)
    for i in range(rows_to_update):
        row = target_table.rows[i + 1]
        cells = row.cells
        record = fortress_records[i]
        
        # 更新各列
        if 'start_time' in col_map and col_map['start_time'] < len(cells):
            clear_and_set_cell(cells[col_map['start_time']], record['start_time'])
        if 'end_time' in col_map and col_map['end_time'] < len(cells):
            clear_and_set_cell(cells[col_map['end_time']], record['end_time'])
        if 'duration' in col_map and col_map['duration'] < len(cells):
            clear_and_set_cell(cells[col_map['duration']], record['duration'])
        if 'source' in col_map and col_map['source'] < len(cells):
            clear_and_set_cell(cells[col_map['source']], record['source'])
        if 'asset_name' in col_map and col_map['asset_name'] < len(cells):
            clear_and_set_cell(cells[col_map['asset_name']], record['asset_name'])
        if 'asset_ip' in col_map and col_map['asset_ip'] < len(cells):
            clear_and_set_cell(cells[col_map['asset_ip']], record['asset_ip'])
        if 'protocol' in col_map and col_map['protocol'] < len(cells):
            clear_and_set_cell(cells[col_map['protocol']], record['protocol'])
        if 'client_type' in col_map and col_map['client_type'] < len(cells):
            clear_and_set_cell(cells[col_map['client_type']], record['client_type'])
    
    # 添加新行
    if len(fortress_records) > existing_rows:
        for i in range(existing_rows, len(fortress_records)):
            record = fortress_records[i]
            new_row = target_table.add_row()
            format_table_row(new_row)  # 确保格式一致
            cells = new_row.cells
            
            if 'start_time' in col_map and col_map['start_time'] < len(cells):
                clear_and_set_cell(cells[col_map['start_time']], record['start_time'])
            if 'end_time' in col_map and col_map['end_time'] < len(cells):
                clear_and_set_cell(cells[col_map['end_time']], record['end_time'])
            if 'duration' in col_map and col_map['duration'] < len(cells):
                clear_and_set_cell(cells[col_map['duration']], record['duration'])
            if 'source' in col_map and col_map['source'] < len(cells):
                clear_and_set_cell(cells[col_map['source']], record['source'])
            if 'asset_name' in col_map and col_map['asset_name'] < len(cells):
                clear_and_set_cell(cells[col_map['asset_name']], record['asset_name'])
            if 'asset_ip' in col_map and col_map['asset_ip'] < len(cells):
                clear_and_set_cell(cells[col_map['asset_ip']], record['asset_ip'])
            if 'protocol' in col_map and col_map['protocol'] < len(cells):
                clear_and_set_cell(cells[col_map['protocol']], record['protocol'])
            if 'client_type' in col_map and col_map['client_type'] < len(cells):
                clear_and_set_cell(cells[col_map['client_type']], record['client_type'])
    
    # 删除多余行
    elif len(fortress_records) < existing_rows:
        rows_to_delete = existing_rows - len(fortress_records)
        for _ in range(rows_to_delete):
            if len(target_table.rows) > 2:
                tbl = target_table._tbl
                tr = target_table.rows[-1]._tr
                tbl.remove(tr)
    
    
    # 更新 VPN 审计记录
    update_vpn_audit(doc, fortress_records)

def update_vpn_audit(doc, fortress_records):
    """更新 VPN 审计记录"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 1. 从已匹配的堡垒机记录中提取用户名，并去重
    vpn_users = set()
    for record in fortress_records:
        user = record.get('user', '')
        if user and user.strip():
            vpn_users.add(user.strip())
    
    if not vpn_users:
        print("   ⚠️ 未从堡垒机提取到用户名，跳过VPN审计更新")
        return
    
    # 2. 读取VPN审计数据并按用户名匹配
    vpn_records = read_vpn_data_by_users(VPN_FILE, vpn_users)
    
    if not vpn_records:
        print("   ⚠️ 未找到匹配的VPN审计记录")
        return
    
    # 3. 查找VPN审计记录表格
    target_table = None
    for i, table in enumerate(doc.tables):
        if len(table.rows) < 2 or len(table.columns) < 6:
            continue
        
        header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if '用户名' in header_text and '用户组' in header_text and '行为' in header_text:
            target_table = table
            break
    
    if not target_table:
        print("   ⚠️ 未找到VPN审计记录表格")
        return
    
    # 4. 分析表头
    headers = [cell.text.strip() for cell in target_table.rows[0].cells]
    col_map = {}
    for idx, h in enumerate(headers):
        if '用户名' in h:
            col_map['username'] = idx
        elif '用户组' in h:
            col_map['user_group'] = idx
        elif '主机IP' in h or '主机ip' in h:
            col_map['host_ip'] = idx
        elif '资源IP' in h or '资源ip' in h:
            col_map['resource_ip'] = idx
        elif '行为' in h:
            col_map['action'] = idx
        elif '时间' in h:
            col_map['time'] = idx
    
    # 5. 更新表格数据
    existing_rows = len(target_table.rows) - 1
    
    # 更新现有行
    rows_to_update = min(len(vpn_records), existing_rows)
    for i in range(rows_to_update):
        row = target_table.rows[i + 1]
        cells = row.cells
        record = vpn_records[i]
        
        if 'username' in col_map and col_map['username'] < len(cells):
            clear_and_set_cell(cells[col_map['username']], record['username'])
        if 'user_group' in col_map and col_map['user_group'] < len(cells):
            clear_and_set_cell(cells[col_map['user_group']], record['user_group'])
        if 'host_ip' in col_map and col_map['host_ip'] < len(cells):
            clear_and_set_cell(cells[col_map['host_ip']], record['host_ip'])
        if 'resource_ip' in col_map and col_map['resource_ip'] < len(cells):
            clear_and_set_cell(cells[col_map['resource_ip']], record['resource_ip'])
        if 'action' in col_map and col_map['action'] < len(cells):
            clear_and_set_cell(cells[col_map['action']], record['action'])
        if 'time' in col_map and col_map['time'] < len(cells):
            clear_and_set_cell(cells[col_map['time']], record['time'])
    
    # 添加新行
    if len(vpn_records) > existing_rows:
        for i in range(existing_rows, len(vpn_records)):
            record = vpn_records[i]
            new_row = target_table.add_row()
            format_table_row(new_row)  # 确保格式一致
            cells = new_row.cells
            
            if 'username' in col_map and col_map['username'] < len(cells):
                clear_and_set_cell(cells[col_map['username']], record['username'])
            if 'user_group' in col_map and col_map['user_group'] < len(cells):
                clear_and_set_cell(cells[col_map['user_group']], record['user_group'])
            if 'host_ip' in col_map and col_map['host_ip'] < len(cells):
                clear_and_set_cell(cells[col_map['host_ip']], record['host_ip'])
            if 'resource_ip' in col_map and col_map['resource_ip'] < len(cells):
                clear_and_set_cell(cells[col_map['resource_ip']], record['resource_ip'])
            if 'action' in col_map and col_map['action'] < len(cells):
                clear_and_set_cell(cells[col_map['action']], record['action'])
            if 'time' in col_map and col_map['time'] < len(cells):
                clear_and_set_cell(cells[col_map['time']], record['time'])
    
    # 删除多余行
    elif len(vpn_records) < existing_rows:
        rows_to_delete = existing_rows - len(vpn_records)
        for _ in range(rows_to_delete):
            if len(target_table.rows) > 2:
                tbl = target_table._tbl
                tr = target_table.rows[-1]._tr
                tbl.remove(tr)
    
def update_snapshot_backup(doc, resource_data, target_year=2026, target_month=3):
    """
    更新快照备份服务报告（附件2）
    
    备份时间：目标月份每周五晚上22:00
    备份负责人：张昊
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 1. 计算目标月份的周五日期
    fridays = get_fridays_from_month(target_year, target_month)
    
    # 2. 从resource_data收集所有主机
    all_hosts = []
    for system_name in sorted(resource_data.keys()):
        data = resource_data[system_name]
        for host in data['hosts']:
            all_hosts.append({
                'host_name': host['host_name'],
                'ip': host['ip']
            })
    
    # 3. 生成备份记录
    backup_records = generate_backup_records(all_hosts, target_year, target_month, backup_person="张昊")
    print(f"   生成 {len(backup_records)} 条备份记录（{len(all_hosts)}台×{len(fridays)}个周五）")
    
    # 4. 查找快照备份表格
    target_table = None
    for i, table in enumerate(doc.tables):
        if len(table.columns) >= 6:
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            if '云主机名称' in header_text and ('备份时间' in header_text or '备份类型' in header_text):
                target_table = table
                break
        elif len(table.columns) >= 5:
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            if '序号' in header_text and '主机IP' in header_text:
                target_table = table
                break
    
    if not target_table:
        print("   ⚠️ 未找到快照备份服务报告表格")
        return
    
    # 5. 分析表头
    headers = [cell.text.strip() for cell in target_table.rows[0].cells]
    col_map = {}
    for idx, h in enumerate(headers):
        if '序号' in h:
            col_map['seq'] = idx
        elif '云主机名称' in h or '主机名称' in h:
            col_map['host_name'] = idx
        elif '主机IP' in h or 'IP' in h:
            col_map['ip'] = idx
        elif '备份时间' in h:
            col_map['backup_time'] = idx
        elif '备份类型' in h:
            col_map['backup_type'] = idx
        elif '备份负责人' in h or '负责人' in h:
            col_map['person'] = idx
    
    # 6. 更新表格数据
    existing_rows = len(target_table.rows) - 1
    
    # 更新现有行
    rows_to_update = min(len(backup_records), existing_rows)
    for i in range(rows_to_update):
        row = target_table.rows[i + 1]
        cells = row.cells
        record = backup_records[i]
        
        if 'seq' in col_map and col_map['seq'] < len(cells):
            clear_and_set_cell(cells[col_map['seq']], str(record['seq']))
        if 'host_name' in col_map and col_map['host_name'] < len(cells):
            clear_and_set_cell(cells[col_map['host_name']], record['host_name'])
        if 'ip' in col_map and col_map['ip'] < len(cells):
            clear_and_set_cell(cells[col_map['ip']], record['ip'])
        if 'backup_time' in col_map and col_map['backup_time'] < len(cells):
            clear_and_set_cell(cells[col_map['backup_time']], record['backup_time'])
        if 'backup_type' in col_map and col_map['backup_type'] < len(cells):
            clear_and_set_cell(cells[col_map['backup_type']], record['backup_type'])
        if 'person' in col_map and col_map['person'] < len(cells):
            clear_and_set_cell(cells[col_map['person']], record['person'])
    
    # 添加新行
    if len(backup_records) > existing_rows:
        for i in range(existing_rows, len(backup_records)):
            record = backup_records[i]
            new_row = target_table.add_row()
            format_table_row(new_row)  # 确保格式一致
            cells = new_row.cells
            
            if 'seq' in col_map and col_map['seq'] < len(cells):
                clear_and_set_cell(cells[col_map['seq']], str(record['seq']))
            if 'host_name' in col_map and col_map['host_name'] < len(cells):
                clear_and_set_cell(cells[col_map['host_name']], record['host_name'])
            if 'ip' in col_map and col_map['ip'] < len(cells):
                clear_and_set_cell(cells[col_map['ip']], record['ip'])
            if 'backup_time' in col_map and col_map['backup_time'] < len(cells):
                clear_and_set_cell(cells[col_map['backup_time']], record['backup_time'])
            if 'backup_type' in col_map and col_map['backup_type'] < len(cells):
                clear_and_set_cell(cells[col_map['backup_type']], record['backup_type'])
            if 'person' in col_map and col_map['person'] < len(cells):
                clear_and_set_cell(cells[col_map['person']], record['person'])
    
    # 删除多余行
    elif len(backup_records) < existing_rows:
        rows_to_delete = existing_rows - len(backup_records)
        for _ in range(rows_to_delete):
            if len(target_table.rows) > 2:
                tbl = target_table._tbl
                tr = target_table.rows[-1]._tr
                tbl.remove(tr)
    
def update_web_tamper_report(doc, target_year=2026, target_month=3):
    """
    更新网页防篡改服务报告
    
    表头：日期、防篡改系统运行状态、防篡改事件监控
    日期：目标月份的每一天
    防篡改系统运行状态：固定为"防篡改系统运行正常"
    防篡改事件监控：固定为"当日无篡改事件发生"
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from calendar import monthrange
    
    # 1. 生成目标月份的所有日期
    _, days_in_month = monthrange(target_year, target_month)
    
    tamper_records = []
    for day in range(1, days_in_month + 1):
        date_str = f"{target_year}年{target_month}月{day}日"
        tamper_records.append({
            'date': date_str,
            'status': '防篡改系统运行正常',
            'monitor': '当日无篡改事件发生'
        })
    
    print(f"   生成 {len(tamper_records)} 天记录")
    
    # 2. 查找网页防篡改表格
    target_table = None
    for i, table in enumerate(doc.tables):
        if len(table.columns) >= 3:
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            if '日期' in header_text and '防篡改' in header_text:
                target_table = table
                break
    
    if not target_table:
        print("   ⚠️ 未找到网页防篡改服务报告表格")
        return
    
    # 3. 分析表头
    headers = [cell.text.strip() for cell in target_table.rows[0].cells]
    col_map = {}
    for idx, h in enumerate(headers):
        if '日期' in h:
            col_map['date'] = idx
        elif '运行状态' in h or '系统运行' in h:
            col_map['status'] = idx
        elif '事件监控' in h or '监控' in h:
            col_map['monitor'] = idx
    
    # 4. 更新表格数据
    existing_rows = len(target_table.rows) - 1
    
    # 更新现有行
    rows_to_update = min(len(tamper_records), existing_rows)
    for i in range(rows_to_update):
        row = target_table.rows[i + 1]
        cells = row.cells
        record = tamper_records[i]
        
        if 'date' in col_map and col_map['date'] < len(cells):
            clear_and_set_cell(cells[col_map['date']], record['date'])
        if 'status' in col_map and col_map['status'] < len(cells):
            clear_and_set_cell(cells[col_map['status']], record['status'])
        if 'monitor' in col_map and col_map['monitor'] < len(cells):
            clear_and_set_cell(cells[col_map['monitor']], record['monitor'])
    
    # 添加新行
    if len(tamper_records) > existing_rows:
        for i in range(existing_rows, len(tamper_records)):
            record = tamper_records[i]
            new_row = target_table.add_row()
            format_table_row(new_row)  # 确保格式一致
            cells = new_row.cells
            
            if 'date' in col_map and col_map['date'] < len(cells):
                clear_and_set_cell(cells[col_map['date']], record['date'])
            if 'status' in col_map and col_map['status'] < len(cells):
                clear_and_set_cell(cells[col_map['status']], record['status'])
            if 'monitor' in col_map and col_map['monitor'] < len(cells):
                clear_and_set_cell(cells[col_map['monitor']], record['monitor'])
    
    # 删除多余行
    elif len(tamper_records) < existing_rows:
        rows_to_delete = existing_rows - len(tamper_records)
        for _ in range(rows_to_delete):
            if len(target_table.rows) > 2:
                tbl = target_table._tbl
                tr = target_table.rows[-1]._tr
                tbl.remove(tr)
    
def read_vpn_data_by_users(file_path, target_users):
    """读取VPN审计数据，根据用户名列表匹配（支持Excel和CSV）"""
    try:
        # 判断文件类型
        if file_path.endswith('.csv'):
            # 读取CSV
            import csv

            
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                all_rows = list(reader)
            
            # 查找表头
            header_row = None
            headers = []
            for i, row in enumerate(all_rows[:10]):
                if any('用户名' in str(cell) for cell in row):
                    header_row = i
                    headers = row
                    break
            
            data_rows = all_rows[header_row+1:] if header_row is not None else []
            
        elif file_path.endswith('.xls'):
            # 读取旧版XLS（尝试忽略加密错误）
            import xlrd
            try:
                book = xlrd.open_workbook(file_path, formatting_info=False)
            except xlrd.biffh.XLRDError as e:
                if 'encrypted' in str(e).lower():
                    book = xlrd.open_workbook(file_path, ignore_workbook_corruption=True)
                else:
                    raise
            ws = book.sheet_by_index(0)
            
            # 读取表头
            header_row = None
            headers = []
            for i in range(min(10, ws.nrows)):
                row_values = ws.row_values(i)
                if any('用户名' in str(cell) for cell in row_values):
                    header_row = i
                    headers = row_values
                    break
            
            # 返回数据行迭代器
            def xls_data_rows():
                for i in range(header_row + 1, ws.nrows):
                    yield ws.row_values(i)
            
            data_rows = xls_data_rows()
            
        else:
            # 读取Excel (xlsx)
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            ws = wb[wb.sheetnames[0]]
            
            # 读取表头
            rows = list(ws.iter_rows(values_only=True, max_row=10))
            header_row = None
            headers = []
            for i, row in enumerate(rows):
                if any('用户名' in str(cell) for cell in row if cell):
                    header_row = i
                    headers = row
                    break
            
            data_rows = ws.iter_rows(values_only=True, min_row=header_row+2)
        
        if not headers:
            print("   ⚠️ 未找到VPN表头")
            if not file_path.endswith('.csv'):
                wb.close()
            return []
        
        # 找到列索引
        col_map = {}
        for idx, h in enumerate(headers):
            if h is None:
                continue
            h_str = str(h).strip()
            if '用户名' in h_str:
                col_map['username'] = idx
            elif '用户组' in h_str:
                col_map['user_group'] = idx
            elif '主机IP' in h_str or '主机ip' in h_str:
                col_map['host_ip'] = idx
            elif '资源IP' in h_str or '资源ip' in h_str:
                col_map['resource_ip'] = idx
            elif '行为' in h_str:
                col_map['action'] = idx
            elif '时间' in h_str:
                col_map['time'] = idx
        
        print(f"   VPN表列映射: {col_map}")
        print(f"   匹配的目标用户名: {list(target_users)[:10]}... (共{len(target_users)}个)")
        
        # 读取并按用户名匹配
        vpn_records = []
        for row in data_rows:
            if not row or len(row) < 5:
                continue
            
            # 获取用户名
            username = str(row[col_map.get('username', 0)]).strip() if col_map.get('username', 0) < len(row) and row[col_map.get('username', 0)] else ''
            
            if not username:
                continue
            
            # 检查用户名是否在目标列表中
            if username not in target_users:
                continue
            
            user_group = str(row[col_map.get('user_group', 1)]).strip() if col_map.get('user_group', 1) < len(row) and row[col_map.get('user_group', 1)] else ''
            
            # 处理时间格式 - 转换为 YYYY/M/D HH:MM 格式（不补零）
            def format_time_vpn(val):
                if not val:
                    return ''
                time_str = str(val).strip()
                
                # 去掉微秒部分
                if '.' in time_str:
                    time_str = time_str.split('.')[0]
                
                # 尝试解析并重新格式化
                try:
                    from datetime import datetime
                    # 尝试多种输入格式
                    for fmt in ['%Y-%m-%d %H:%M:%S', '%Y/%m/%d %H:%M:%S', '%Y-%m-%d %H:%M', '%Y/%m/%d %H:%M']:
                        try:
                            dt = datetime.strptime(time_str, fmt)
                            # 输出格式: YYYY/M/D HH:MM (不要秒，不补零)
                            return f"{dt.year}/{dt.month}/{dt.day} {dt.hour}:{dt.minute:02d}"
                        except ValueError:
                            continue
                except Exception:
                    pass
                
                # 如果解析失败，进行简单的字符串替换
                # 将 2026-03-31 15:50:25 转换为 2026/3/31 15:50
                time_str = time_str.replace('-', '/')
                if ':' in time_str:
                    parts = time_str.split(':')
                    if len(parts) >= 3:
                        # 去掉秒部分
                        time_str = parts[0] + ':' + parts[1]
                
                return time_str
            
            record = {
                'username': username,
                'user_group': user_group,
                'host_ip': str(row[col_map.get('host_ip', 2)]).strip() if col_map.get('host_ip', 2) < len(row) and row[col_map.get('host_ip', 2)] else '',
                'resource_ip': str(row[col_map.get('resource_ip', 3)]).strip() if col_map.get('resource_ip', 3) < len(row) and row[col_map.get('resource_ip', 3)] else '',
                'action': str(row[col_map.get('action', 4)]).strip() if col_map.get('action', 4) < len(row) and row[col_map.get('action', 4)] else '',
                'time': format_time_vpn(row[col_map.get('time', 5)]) if col_map.get('time', 5) < len(row) else '',
            }
            
            vpn_records.append(record)
        
        # 关闭文件
        if file_path.endswith('.xlsx') and 'wb' in locals():
            wb.close()
        # XLS格式由xlrd自动管理，无需手动关闭
        # CSV格式无需关闭
        
        return vpn_records
        
    except Exception as e:
        print(f"❌ 读取VPN数据失败: {e}")
        import traceback
        traceback.print_exc()
        return []

# ============ 审核功能 ============
def audit_report_dates(output_path, target_year, target_month):
    """
    审核报告中日期的正确性
    确保所有日期都在目标月份范围内
    """
    from docx import Document
    from datetime import datetime
    import re
    
    print("\n📋 步骤8: 审核日期数据...")
    print("-"*60)
    
    try:
        doc = Document(output_path)
        
        # 计算目标月份的开始和结束日期
        target_start = f"{target_year}-{target_month:02d}-01"
        if target_month == 12:
            target_end = f"{target_year}-{target_month:02d}-31"
        else:
            target_end = f"{target_year}-{(target_month+1):02d}-01"
        
        print(f"审核: {target_year}年{target_month}月")
        
        issues = []
        date_pattern = r'(\d{4}-\d{2}-\d{2})'
        
        # 检查所有表格中的日期
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    dates = re.findall(date_pattern, text)
                    
                    for date_str in dates:
                        try:
                            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                            if date_obj.year != target_year or date_obj.month != target_month:
                                issues.append({
                                    'table': table_idx,
                                    'row': row_idx,
                                    'cell': cell_idx,
                                    'date': date_str,
                                    'text': text[:50],
                                    'issue': f"日期不在目标月份内"
                                })
                        except:
                            pass
        
        if issues:
            print(f"\n⚠️ 发现 {len(issues)} 处日期异常")
            for i, issue in enumerate(issues[:5], 1):
                print(f"  {i}. 表格{issue['table']}行{issue['row']}: {issue['date']}")
        else:
            print(" ✅ 日期审核通过")
        
        return len(issues) == 0
        
    except Exception as e:
        print(f"❌ 审核过程出错: {e}")
        return False

# ============ 主程序 ============
def get_default_target_month():
    """获取默认目标月份（上个月）"""
    from datetime import datetime
    now = datetime.now()
    year = now.year
    month = now.month - 1
    if month == 0:
        year -= 1
        month = 12
    return year, month

if __name__ == '__main__':
    print("="*60)
    print("🚀 月报自动化处理系统 v2")
    print("="*60)
    
    # 获取默认目标月份（上个月）
    target_year, target_month = get_default_target_month()
    print(f"\n默认目标月份: {target_year}年{target_month}月（上个月）")
    
    # 确保输出目录存在
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 支持命令行参数选择模板
    import sys
    if len(sys.argv) > 1:
        # 从命令行获取模板文件名
        test_template = sys.argv[1]
    else:
        # 默认处理农林科
        test_template = '政务云服务运维月报-2025年11月-北京市农林科学院.docx'
    
    template_path = os.path.join(TEMPLATE_DIR, test_template)
    from datetime import datetime
    timestamp = datetime.now().strftime('%H%M%S')
    output_name = test_template.replace('2025年11月', f'{target_year}年{target_month:02d}_v2_{timestamp}')
    output_path = os.path.join(OUTPUT_DIR, output_name)
    
    if os.path.exists(template_path):
        success = generate_report(template_path, output_path)
        if success:
            print(f"\n✅ 成功生成: {output_path}")
            # 审核日期
            audit_report_dates(output_path, target_year, target_month)
        else:
            print(f"\n❌ 生成失败")
    else:
        print(f"❌ 模板不存在: {template_path}")
        print(f"可用模板（部分）:")
        templates = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith('.docx') and not f.startswith('~$')]
        for t in templates[:10]:
            print(f"  - {t}")
    
    print("\n" + "="*60)
    print("处理完成")
    print("="*60)
