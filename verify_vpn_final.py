#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 找到最新农林科文件
output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '农林' in f and f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
output_file = os.path.join(output_dir, files[0])

print(f"验证文件: {files[0]}\n")
print("="*70)
print("✅ 验证 VPN 审计记录")
print("="*70)

doc = Document(output_file)

# 查找 VPN 审计记录表
vpn_table = None
for i, table in enumerate(doc.tables):
    if len(table.columns) >= 6:
        header = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if '用户名' in header and '用户组' in header and '行为' in header:
            vpn_table = table
            print(f"\n找到 VPN 审计记录表 (表格{i}):")
            print(f"  列数: {len(table.columns)}, 行数: {len(table.rows)}")
            break

if vpn_table:
    # 显示表头
    headers = [cell.text.strip() for cell in vpn_table.rows[0].cells]
    print(f"\n表头: {headers}")
    
    # 显示前10行数据
    print("\n前10行数据:")
    print("-"*70)
    for row_idx in range(1, min(11, len(vpn_table.rows))):
        cells = [cell.text.strip() for cell in vpn_table.rows[row_idx].cells]
        print(f"  行{row_idx}: {cells}")
    
    # 统计
    print("\n" + "="*70)
    print(f"✅ VPN 审计记录统计:")
    print(f"  总记录数: {len(vpn_table.rows) - 1} 条")
    
    # 验证格式（居中）
    if len(vpn_table.rows) > 1:
        sample_cell = vpn_table.rows[1].cells[0]
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        alignment = sample_cell.paragraphs[0].alignment if sample_cell.paragraphs else None
        is_centered = alignment == WD_ALIGN_PARAGRAPH.CENTER
        print(f"  数据是否居中: {'是 ✓' if is_centered else '否 ✗'}")
    
    print("="*70)
else:
    print("❌ 未找到 VPN 审计记录表")
