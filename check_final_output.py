#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 找到最新的输出文件
output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '健康' in f and f.endswith('.docx') and not f.startswith('~$')]
if not files:
    print("未找到文件")
    sys.exit(1)

files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
latest_file = os.path.join(output_dir, files[0])

print(f"检查文件: {files[0]}\n")

doc = Document(latest_file)

print(f"文档中共有 {len(doc.tables)} 个表格\n")

for i, table in enumerate(doc.tables[:10]):
    # 获取表头
    headers = []
    for row in table.rows[:2]:
        row_text = ' '.join([c.text.strip() for c in row.cells])
        if row_text:
            headers.append(row_text[:60])
    
    print(f"表格{i}: {len(table.rows)}行 x {len(table.columns)}列")
    print(f"  表头: {headers}")
    print()
