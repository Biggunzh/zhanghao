#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证修复后的表格数据"""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_175234-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 验证修复后的表格数据")
    print("="*70)
    
    # 检查CPU使用率表
    print("\n📊 CPU使用率汇总表:")
    table1 = doc.tables[1]
    for row_idx, row in enumerate(table1.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    # 检查存储使用率表
    print("\n📊 存储使用率汇总表:")
    table3 = doc.tables[3]
    for row_idx, row in enumerate(table3.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    # 验证特定值
    print("\n" + "="*70)
    print("✅ 数据验证:")
    print("="*70)
    
    # CPU表第4行（北京市卫生人事代理系统）
    cpu_row = table1.rows[4].cells
    cpu_total = cpu_row[3].text.strip()
    print(f"北京市卫生人事代理系统 CPU总量: {cpu_total} (应为16) {'✓' if cpu_total == '16' else '✗'}")
    
    # 存储表各行的值
    storage_data = [
        (2, "北京市卫生高级技术职务申报和评审系统", "3550"),
        (3, "北京市卫生考试考务系统", "2150"),
        (4, "北京市卫生人事代理系统", "3100"),
        (5, "北京市卫生系统事业单位招聘网上申报系统", "1200"),
        (6, "北京卫生人才网", "13350"),
    ]
    
    print("\n存储总量验证:")
    for row_idx, system_name, expected in storage_data:
        storage_row = table3.rows[row_idx].cells
        storage_total = storage_row[3].text.strip()
        status = '✓' if storage_total == expected else '✗'
        print(f"  {system_name}: {storage_total} (应为{expected}) {status}")
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
