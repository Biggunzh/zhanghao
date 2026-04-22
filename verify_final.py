#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '农林' in f and f.endswith('.docx')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)

if not files:
    print("未找到文件")
    exit()

fp = os.path.join(output_dir, files[0])
doc = Document(fp)

print(f"文件: {files[0]}")
print(f"大小: {os.path.getsize(fp):,} bytes")
print()

# 验证关键表格
for i, table in enumerate(doc.tables):
    headers = [c.text.strip() for c in table.rows[0].cells]
    h = ' '.join(headers)
    
    if '备份类型' in h:
        bt = table.rows[1].cells[headers.index('备份类型')].text
        print(f'表格{i} 快照备份类型: {bt}')
    elif '日期' in h and '防篡改' in h:
        d = table.rows[1].cells[0].text
        print(f'表格{i} 防篡改日期: {d}')
        # 检查居中
        centered = all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER 
                      for row in table.rows[1:4] 
                      for cell in row.cells 
                      for p in cell.paragraphs)
        print(f'      居中: {"是" if centered else "否"}')
    elif '平均使用率' in h or '当前使用率' in h:
        # 附件2表格
        centered = all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER 
                      for row in table.rows[1:4] 
                      for cell in row.cells 
                      for p in cell.paragraphs)
        print(f'表格{i} ({h[:15]}...): {len(table.rows)-1}行, 居中:{"是" if centered else "否"}')

print()
print('农林科2026年3月月报生成完成!')
