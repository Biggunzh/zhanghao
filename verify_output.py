#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月-北京市农林科学院.docx'

try:
    doc = Document(output_file)
    print("✅ 文档打开成功！")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")
    
    # 显示前5段内容
    print("\n前5段内容:")
    for i, para in enumerate(doc.paragraphs[:5]):
        if para.text.strip():
            print(f"{i+1}. {para.text[:80]}")
    
    # 检查日期是否替换成功
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    if '2026年3月' in full_text or '2026年03月' in full_text:
        print("\n✅ 日期替换成功!")
    else:
        print("\n⚠️ 日期可能未替换")
    
    # 显示第一个表格的信息
    if doc.tables:
        print(f"\n第一个表格: {len(doc.tables[0].rows)}行 x {len(doc.tables[0].columns)}列")
        
except Exception as e:
    print(f"❌ 错误: {e}")
