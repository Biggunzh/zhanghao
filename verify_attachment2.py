#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证附件2（资源使用率详情）的居中格式"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
doc = Document(os.path.join(output_dir, files[0]))

print(f"验证文件: {files[0]}\n")

# 查找附件2表格（CPU/内存/磁盘使用率详情）
for i, table in enumerate(doc.tables):
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    header_str = ' '.join(headers)
    
    # 判断是否是附件2类型表格
    if ('CPU' in header_str or '内存' in header_str or '磁盘' in header_str or '存储' in header_str) and \
       ('平均' in header_str or '最大' in header_str):
        
        print(f"表格{i}: {header_str[:40]}...")
        print(f"  总行数: {len(table.rows)} 行")
        
        # 检查前3行的对齐方式
        all_centered = True
        for row_idx in range(1, min(4, len(table.rows))):
            row = table.rows[row_idx]
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        all_centered = False
                        break
        
        status = '是' if all_centered else '否'
        print(f"  前3行全部居中: {status}")
        
        # 显示示例数据
        if len(table.rows) > 1:
            row = table.rows[1]
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  示例数据: {cells[:5]}")
        print()
