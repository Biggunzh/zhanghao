#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证所有表格居中对齐"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)

if not files:
    print("未找到文件")
    exit()

fp = os.path.join(output_dir, files[0])
doc = Document(fp)

print(f"验证文件: {files[0]}")
print(f"文件大小: {os.path.getsize(fp):,} bytes")
print("\n" + "="*60)
print("表格居中对齐验证")
print("="*60)

all_passed = True
for i, table in enumerate(doc.tables):
    headers = [c.text.strip() for c in table.rows[0].cells]
    h = ' '.join(headers[:3])
    
    # 检查前3行数据是否居中
    centered = True
    sample_rows = min(4, len(table.rows))
    for row_idx in range(1, sample_rows):
        for cell in table.rows[row_idx].cells:
            for para in cell.paragraphs:
                if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    centered = False
                    break
        if not centered:
            break
    
    status = "OK" if centered else "FAIL"
    if not centered:
        all_passed = False
    
    # 简化表头显示
    short_name = h[:20] + "..." if len(h) > 20 else h
    print(f"Table {i}: {short_name:25s} {len(table.rows)-1:4d} rows  {status}")

print("="*60)
if all_passed:
    print("All tables are centered!")
else:
    print("Some tables not fully centered")
print("="*60)
