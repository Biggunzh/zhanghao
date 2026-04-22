#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 找到最新司法局文件
output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '司法局' in f and f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
output_file = os.path.join(output_dir, files[0])

print(f"验证文件: {files[0]}\n")

doc = Document(output_file)

print("📋 1. 基础资源台账概况")
print("-"*70)
for para in doc.paragraphs[:30]:
    text = para.text.strip()
    if '共有' in text and '业务系统' in text:
        print(text)
        break

print("\n📋 2. 本月技术支撑统计")
print("-"*70)
table0 = doc.tables[0]
for row in table0.rows:
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  {cells}")

print("\n📋 3. CPU/内存/存储使用率汇总")
print("-"*70)

# CPU表
table1 = doc.tables[1]
print("\nCPU使用率:")
for row_idx in range(1, min(5, len(table1.rows))):
    cells = [cell.text.strip() for cell in table1.rows[row_idx].cells]
    print(f"  {cells}")

# 存储表
table3 = doc.tables[3]
print("\n存储使用率:")
for row_idx in range(1, min(5, len(table3.rows))):
    cells = [cell.text.strip() for cell in table3.rows[row_idx].cells]
    print(f"  {cells}")

print("\n" + "="*70)
print("✅ 数据验证:")
print("="*70)

# 验证CPU
cpu_expected = [
    ("北京市规范性文件审查工作信息平台", "10"),
    ("北京市行政执法信息服务平台", "534"),
    ("政务综合管理平台", "16"),
]

print("\nCPU总量验证:")
for row_idx in range(2, 5):
    name = table1.rows[row_idx].cells[1].text.strip()
    actual = table1.rows[row_idx].cells[3].text.strip()
    expected = cpu_expected[row_idx-2][1]
    status = '✓' if actual == expected else '✗'
    print(f"  {name}: {actual} (应为{expected}) {status}")

# 验证存储
storage_expected = [
    ("北京市规范性文件审查工作信息平台", "3600"),
    ("北京市行政执法信息服务平台", "103308"),
    ("政务综合管理平台", "2100"),
]

print("\n存储总量验证:")
for row_idx in range(2, 5):
    name = table3.rows[row_idx].cells[1].text.strip()
    actual = table3.rows[row_idx].cells[3].text.strip()
    expected = storage_expected[row_idx-2][1]
    status = '✓' if actual == expected else '✗'
    print(f"  {name}: {actual} (应为{expected}) {status}")

print("\n" + "="*70)
print("✅ 司法局报告验证完成！")
print("="*70)
