#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月-北京市农林科学院.docx'

try:
    doc = Document(output_file)
    print("✅ 文档打开成功！")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")
    
    # 查找包含"2026年3月"的段落
    print("\n包含'2026年3月'的关键段落:")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '2026年3月' in text and '共有' in text and '运行' in text:
            print(f"\n段落 {i}:")
            print(text)
            print(f"\n文本长度: {len(text)}")
            # 检查重复
            repeat_cnt = text.count('北京市农林科学院')
            print(f"'北京市农林科学院'出现次数: {repeat_cnt}")
            if repeat_cnt > 1:
                print("⚠️ 警告: 文本有重复!")
            else:
                print("✅ 文本正常，无重复")
            break
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
