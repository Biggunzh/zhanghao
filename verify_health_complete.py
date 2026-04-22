#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证北京市卫生健康人力资源发展中心完整输出"""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_173106-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 北京市卫生健康人力资源发展中心 - 完整验证")
    print("="*70)
    
    # 1. 基础资源台账概况
    print("\n📋 1. 基础资源台账概况")
    print("-"*70)
    for para in doc.paragraphs[:50]:
        text = para.text.strip()
        if '共有' in text and '业务系统' in text and '运行' in text:
            print(text[:100] + "...")
            break
    
    # 2. 统计各部分内容
    print("\n📊 2. 表格统计")
    print("-"*70)
    
    table_summary = []
    for i, table in enumerate(doc.tables[:10]):
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        header_str = ' '.join(headers[:3]) if headers else '未知'
        table_summary.append((i, len(table.rows), len(table.columns), header_str[:30]))
    
    for idx, rows, cols, header in table_summary:
        print(f"  表{idx}: {rows}行 x {cols}列 - {header}")
    
    # 3. 本月技术支撑统计
    print("\n📋 3. 本月技术支撑统计")
    print("-"*70)
    table0 = doc.tables[0]
    for row in table0.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    # 4. 资源使用率汇总
    print("\n📋 4. CPU/内存/存储使用率汇总表")
    print("-"*70)
    for i in range(1, 4):
        if i < len(doc.tables):
            table = doc.tables[i]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"\n  表{i} 表头: {headers[:5]}")
            # 显示数据行（前3行）
            for row_idx in range(1, min(4, len(table.rows))):
                cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
                print(f"    {cells[:5]}")
    
    # 5. 附件1
    print("\n📋 5. 附件1 - 基础资源台账")
    print("-"*70)
    if len(doc.tables) > 4:
        table4 = doc.tables[4]
        headers = [cell.text.strip() for cell in table4.rows[0].cells]
        print(f"  表4 表头: {headers}")
        print(f"  总行数: {len(table4.rows)}行")
        print("  前3行:")
        for row_idx in range(1, min(4, len(table4.rows))):
            cells = [cell.text.strip() for cell in table4.rows[row_idx].cells]
            print(f"    {cells[:6]}")
    
    # 6. 附件2
    print("\n📋 6. 附件2 - 资源使用率详情")
    print("-"*70)
    for i in range(5, 8):
        if i < len(doc.tables):
            table = doc.tables[i]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            header_str = ' '.join(headers[:3])
            print(f"  表{i}: {header_str}... ({len(table.rows)}行)")
    
    print("\n" + "="*70)
    print("✅ 验证完成！")
    print("="*70)
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
