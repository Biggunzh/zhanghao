#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 找到最新文件
output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '健康' in f and f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
output_file = os.path.join(output_dir, files[0])

print(f"验证文件: {files[0]}\n")

doc = Document(output_file)
table1 = doc.tables[1]  # CPU
table3 = doc.tables[3]  # 存储

print("📊 CPU使用率汇总表:")
for row_idx, row in enumerate(table1.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  行{row_idx}: {cells}")

print("\n📊 存储使用率汇总表:")
for row_idx, row in enumerate(table3.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  行{row_idx}: {cells}")

print("\n" + "="*70)
print("✅ 数据验证:")

# 验证CPU
cpu_checks = [
    (2, "北京市卫生高级技术职务申报和评审系统", "104"),
    (3, "北京市卫生考试考务系统", "32"),
    (4, "北京市卫生人事代理系统", "16"),
    (5, "北京市卫生系统事业单位招聘网上申报系统", "64"),
    (6, "北京卫生人才网", "80"),
]

all_pass = True
for row_idx, name, expected in cpu_checks:
    actual = table1.rows[row_idx].cells[3].text.strip()
    status = '✓' if actual == expected else '✗'
    if actual != expected:
        all_pass = False
    print(f"  {name} CPU: {actual} (应为{expected}) {status}")

# 验证存储
storage_checks = [
    (2, "北京市卫生高级技术职务申报和评审系统", "3550"),
    (3, "北京市卫生考试考务系统", "2150"),
    (4, "北京市卫生人事代理系统", "3100"),
    (5, "北京市卫生系统事业单位招聘网上申报系统", "1200"),
    (6, "北京卫生人才网", "13350"),
]

for row_idx, name, expected in storage_checks:
    actual = table3.rows[row_idx].cells[3].text.strip()
    status = '✓' if actual == expected else '✗'
    if actual != expected:
        all_pass = False
    print(f"  {name} 存储: {actual} (应为{expected}) {status}")

if all_pass:
    print("\n✅✅✅ 全部修复成功！所有数据正确！")
else:
    print("\n❌ 仍有问题")
