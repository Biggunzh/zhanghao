#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_180233-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 验证修正后的数据")
    print("="*70)
    
    table1 = doc.tables[1]  # CPU
    table3 = doc.tables[3]  # 存储
    
    print("\n📊 CPU使用率汇总表:")
    for row_idx, row in enumerate(table1.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行{row_idx}: {cells}")
    
    print("\n📊 存储使用率汇总表:")
    for row_idx, row in enumerate(table3.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行{row_idx}: {cells}")
    
    print("\n" + "="*70)
    print("✅ 关键数据验证:")
    print("="*70)
    
    # 验证CPU
    print("\nCPU总量验证:")
    cpu_checks = [
        (2, "北京市卫生高级技术职务申报和评审系统", "104"),
        (3, "北京市卫生考试考务系统", "32"),
        (4, "北京市卫生人事代理系统", "16"),
        (5, "北京市卫生系统事业单位招聘网上申报系统", "64"),
        (6, "北京卫生人才网", "80"),
    ]
    
    all_pass = True
    for row_idx, name, expected in cpu_checks:
        actual = table1.rows[row_idx].cells[3].text.strip()
        status = '✓' if actual == expected else '✗'
        if actual != expected:
            all_pass = False
        print(f"  {name}: {actual} (应为{expected}) {status}")
    
    # 验证存储
    print("\n存储总量验证:")
    storage_checks = [
        (2, "北京市卫生高级技术职务申报和评审系统", "3550"),
        (3, "北京市卫生考试考务系统", "2150"),
        (4, "北京市卫生人事代理系统", "3100"),
        (5, "北京市卫生系统事业单位招聘网上申报系统", "1200"),
        (6, "北京卫生人才网", "13350"),
    ]
    
    for row_idx, name, expected in storage_checks:
        actual = table3.rows[row_idx].cells[3].text.strip()
        status = '✓' if actual == expected else '✗'
        if actual != expected:
            all_pass = False
        print(f"  {name}: {actual} (应为{expected}) {status}")
    
    print("\n" + "="*70)
    if all_pass:
        print("✅✅✅ 全部修复成功！所有数据正确！")
    else:
        print("❌ 仍有问题")
    print("="*70)
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
