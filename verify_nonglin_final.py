#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 找到最新农林科文件
output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '农林' in f and f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
output_file = os.path.join(output_dir, files[0])

print(f"验证文件: {files[0]}\n")
print("="*70)
print("✅ 北京市农林科学院 - 完整报告验证")
print("="*70)

doc = Document(output_file)

# 1. 基础资源台账
print("\n📋 1. 基础资源台账概况")
print("-"*70)
for para in doc.paragraphs[:30]:
    text = para.text.strip()
    if '共有' in text and '业务系统' in text and '运行' in text:
        print(text)
        break

# 2. 本月技术支撑统计
print("\n📋 2. 本月技术支撑统计")
print("-"*70)
table0 = doc.tables[0]
for row in table0.rows:
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  {cells}")

# 3. CPU/内存/存储使用率汇总
print("\n📋 3. 资源使用率汇总表")
print("-"*70)
for i, title in [(1, 'CPU'), (2, '内存'), (3, '存储')]:
    table = doc.tables[i]
    print(f"\n{title}使用率:")
    for row_idx in range(min(4, len(table.rows))):
        cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
        print(f"  {cells}")

# 4. 附件1
print("\n📋 4. 附件1 - 基础资源台账")
print("-"*70)
table4 = doc.tables[4]
print(f"  总行数: {len(table4.rows)}行")
print("  前5行:")
for row_idx in range(1, min(6, len(table4.rows))):
    cells = [cell.text.strip() for cell in table4.rows[row_idx].cells]
    print(f"    {cells[:5]}")

# 5. 附件3
print("\n📋 5. 附件3 - 堡垒机审计记录")
print("-"*70)
fortress_table = None
for table in doc.tables:
    if len(table.columns) >= 8:
        header = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if '开始时间' in header and '资产IP' in header:
            fortress_table = table
            break

if fortress_table:
    print(f"  总行数: {len(fortress_table.rows)}行")
    print("  前3行:")
    for row_idx in range(1, min(4, len(fortress_table.rows))):
        cells = [cell.text.strip() for cell in fortress_table.rows[row_idx].cells]
        print(f"    {cells[:4]}...")

# 验证数据
print("\n" + "="*70)
print("✅ 数据验证:")
print("="*70)

# 验证总量
table1 = doc.tables[1]
table3 = doc.tables[3]

cpu_data = [
    ("微营销", "61"),
    ("长城网", "838"),
]

storage_data = [
    ("微营销", "3850"),
    ("长城网", "71796"),
]

print("\nCPU总量:")
all_pass = True
for row_idx in range(2, 4):
    name = table1.rows[row_idx].cells[1].text.strip()
    actual = table1.rows[row_idx].cells[3].text.strip()
    expected = cpu_data[row_idx-2][1]
    status = '✓' if actual == expected else '✗'
    if actual != expected:
        all_pass = False
    print(f"  {name}: {actual} (应为{expected}) {status}")

print("\n存储总量:")
for row_idx in range(2, 4):
    name = table3.rows[row_idx].cells[1].text.strip()
    actual = table3.rows[row_idx].cells[3].text.strip()
    expected = storage_data[row_idx-2][1]
    status = '✓' if actual == expected else '✗'
    if actual != expected:
        all_pass = False
    print(f"  {name}: {actual} (应为{expected}) {status}")

if all_pass:
    print("\n✅✅✅ 全部验证通过！")
else:
    print("\n❌ 有数据错误")

print("="*70)
