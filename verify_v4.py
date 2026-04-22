#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_175843-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 验证 cell.text 方式修复")
    print("="*70)
    
    table1 = doc.tables[1]  # CPU
    table3 = doc.tables[3]  # 存储
    
    print("\n📊 CPU使用率汇总表:")
    for row in table1.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    print("\n📊 存储使用率汇总表:")
    for row in table3.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    print("\n✅ 关键数据验证:")
    
    # 验证CPU表
    row_idx = 4  # 北京市卫生人事代理系统
    cpu_row = table1.rows[row_idx].cells
    cpu_total = cpu_row[3].text.strip()
    print(f"\n北京市卫生人事代理系统 CPU: {cpu_total} (应为16) {'✓' if cpu_total == '16' else '✗'}")
    
    # 验证存储表
    storage_checks = [
        (2, "北京市卫生高级技术职务申报和评审系统", "3550"),
        (3, "北京市卫生人事代理系统", "3100"),
        (6, "北京卫生人才网", "13350"),
    ]
    
    all_pass = True
    for row_idx, system_name, expected in storage_checks:
        storage_row = table3.rows[row_idx].cells
        storage_total = storage_row[3].text.strip()
        status = '✓' if storage_total == expected else '✗'
        if storage_total != expected:
            all_pass = False
        print(f"{system_name} 存储: {storage_total} (应为{expected}) {status}")
    
    if all_pass and cpu_total == '16':
        print("\n✅✅✅ 修复成功！所有数据正确！")
    else:
        print("\n❌ 仍有问题")
    
except Exception as e:
    print(f"❌ 错误: {e}")
