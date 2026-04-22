#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_175724-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 验证XML方式修复后的数据")
    print("="*70)
    
    table1 = doc.tables[1]  # CPU
    table3 = doc.tables[3]  # 存储
    
    print("\n📊 CPU使用率汇总表:")
    for row in table1.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    print("\n📊 存储使用率汇总表:")
    for row in table3.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    print("\n✅ 数据验证:")
    
    # 验证CPU表
    cpu_data = [
        (4, "北京市卫生人事代理系统", "16"),
    ]
    
    print("\n关键数据验证:")
    all_pass = True
    
    row_idx, system_name, expected = cpu_data[0]
    cpu_row = table1.rows[row_idx].cells
    cpu_total = cpu_row[3].text.strip()
    status = '✓' if cpu_total == expected else '✗'
    if cpu_total != expected:
        all_pass = False
    print(f"  {system_name} CPU: {cpu_total} (应为{expected}) {status}")
    
    # 验证存储表关键数据
    storage_checks = [
        (2, "北京市卫生高级技术职务申报和评审系统", "3550"),
        (4, "北京市卫生人事代理系统", "3100"),
        (6, "北京卫生人才网", "13350"),
    ]
    
    for row_idx, system_name, expected in storage_checks:
        storage_row = table3.rows[row_idx].cells
        storage_total = storage_row[3].text.strip()
        status = '✓' if storage_total == expected else '✗'
        if storage_total != expected:
            all_pass = False
        print(f"  {system_name} 存储: {storage_total} (应为{expected}) {status}")
    
    if all_pass:
        print("\n✅ 修复成功！")
    else:
        print("\n❌ 仍有问题")
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
