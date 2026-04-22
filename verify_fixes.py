#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os

output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
doc = Document(os.path.join(output_dir, files[0]))

print(f"验证文件: {files[0]}\n")

# 1. 验证快照备份类型
print("="*60)
print("1. 快照备份服务报告 - 备份类型验证")
print("="*60)
for i, table in enumerate(doc.tables):
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    if '备份类型' in headers:
        print(f"\n表格{i}:")
        # 检查前3行的备份类型
        for row_idx in range(1, min(4, len(table.rows))):
            backup_type = table.rows[row_idx].cells[headers.index('备份类型')].text.strip()
            print(f"  行{row_idx}: {backup_type}")
        break

# 2. 验证网页防篡改日期格式
print("\n" + "="*60)
print("2. 网页防篡改服务报告 - 日期格式验证")
print("="*60)
for i, table in enumerate(doc.tables):
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    if '日期' in headers and '防篡改' in headers:
        print(f"\n表格{i}:")
        # 检查前5行和最后1行的日期格式
        for row_idx in [1, 2, 3, 4, 5, len(table.rows)-1]:
            if row_idx < len(table.rows):
                date_str = table.rows[row_idx].cells[headers.index('日期')].text.strip()
                print(f"  行{row_idx}: {date_str}")
        break

print("\n" + "="*60)
print("验证完成!")
print("="*60)
