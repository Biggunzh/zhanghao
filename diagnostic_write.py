#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""诊断表格写入问题"""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

template = r'D:\月报自动化\月报模板\政务云服务运维月报-2025年11月-北京市卫生健康人力资源发展中心.docx'
doc = Document(template)

table1 = doc.tables[1]  # CPU表

print("="*70)
print("诊断表格写入")
print("="*70)

# 获取第4行（北京市卫生人事代理系统）
row_idx = 4
cell = table1.rows[row_idx].cells[3]  # CPU总量列

print(f"\n行{row_idx}的CPU总量列:")
print(f"  写入前值: '{cell.text}'")

# 写入新值
from docx.enum.text import WD_ALIGN_PARAGRAPH
cell.text = "16"

print(f"  写入后值: '{cell.text}'")

# 检查段落数
print(f"  段落数: {len(cell.paragraphs)}")
for i, para in enumerate(cell.paragraphs):
    print(f"    段落{i}: '{para.text}' (runs: {len(para.runs)})")
    for j, run in enumerate(para.runs):
        print(f"      run{j}: '{run.text}'")

# 保存到新文件测试
test_output = r'D:\月报自动化\输出月报\test_write.docx'
doc.save(test_output)
print(f"\n  已保存到: {test_output}")

# 重新打开验证
doc2 = Document(test_output)
cell2 = doc2.tables[1].rows[4].cells[3]
print(f"  重新打开后值: '{cell2.text}'")
