#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '农林' in f and f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
output_file = os.path.join(output_dir, files[0])

print(f"验证文件: {files[0]}")
print(f"文件大小: {os.path.getsize(output_file):,} bytes\n")
print("="*70)
print("✅ 北京市农林科学院 - 标准月报验证")
print("="*70)

doc = Document(output_file)

# 1. 基础资源台账概况
print("\n📋 1. 基础资源台账概况")
print("-"*70)
for para in doc.paragraphs[:30]:
    text = para.text.strip()
    if '共有' in text and '业务系统' in text:
        print(text)
        break

# 2. 本月技术支撑统计
print("\n📋 2. 本月技术支撑统计")
print("-"*70)
table0 = doc.tables[0]
for row in table0.rows:
    cells = [cell.text.strip() for cell in row.cells]
    print(f"  {cells}")

# 3. 资源使用率汇总
print("\n📋 3. 资源使用率汇总表")
print("-"*70)

for i, title in [(1, 'CPU'), (2, '内存'), (3, '存储')]:
    table = doc.tables[i]
    print(f"\n{title}使用率:")
    for row_idx in range(min(6, len(table.rows))):
        cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
        print(f"  {cells}")

# 4. 附件1 - 基础资源台账
print("\n📋 4. 附件1 - 基础资源台账")
print("-"*70)
table4 = doc.tables[4]
print(f"  总行数: {len(table4.rows)}行")
print("  前3行:")
for row_idx in range(1, min(4, len(table4.rows))):
    cells = [cell.text.strip() for cell in table4.rows[row_idx].cells]
    print(f"    {cells[:4]}")

# 5. 附件3 - 堡垒机审计
print("\n📋 5. 附件3 - 堡垒机审计记录")
print("-"*70)
fortress_table = None
for table in doc.tables:
    if len(table.columns) >= 8:
        header = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if '开始时间' in header and '资产IP' in header:
            fortress_table = table
            break

if fortress_table:
    print(f"  总行数: {len(fortress_table.rows)}行")
    print("  前3行:")
    for row_idx in range(1, min(4, len(fortress_table.rows))):
        cells = [cell.text.strip() for cell in fortress_table.rows[row_idx].cells]
        print(f"    开始={cells[0]}, 资产={cells[5]}, IP={cells[6]}")

# 6. 附件3 - VPN审计
print("\n📋 6. 附件3 - VPN审计记录")
print("-"*70)
vpn_table = None
for table in doc.tables:
    if len(table.columns) >= 6:
        header = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if '用户名' in header and '用户组' in header:
            vpn_table = table
            break

if vpn_table:
    print(f"  总行数: {len(vpn_table.rows)}行")
    
    # 统计用户组
    user_groups = {}
    for row_idx in range(1, len(vpn_table.rows)):
        user_group = vpn_table.rows[row_idx].cells[1].text.strip()
        user_groups[user_group] = user_groups.get(user_group, 0) + 1
    
    print(f"\n  用户组分布:")
    for group, count in sorted(user_groups.items(), key=lambda x: x[1], reverse=True):
        print(f"    {group}: {count} 条")
    
    print(f"\n  前3行:")
    for row_idx in range(1, min(4, len(vpn_table.rows))):
        cells = [cell.text.strip() for cell in vpn_table.rows[row_idx].cells]
        print(f"    {cells}")

# 验证关键数据
print("\n" + "="*70)
print("✅ 关键数据验证")
print("="*70)

# 检查业务系统
table1 = doc.tables[1]
systems = []
for row_idx in range(1, len(table1.rows)):
    name = table1.rows[row_idx].cells[1].text.strip()
    if name and name != '合计':
        systems.append(name)

print(f"\n业务系统 ({len(systems)} 个):")
for s in systems:
    print(f"  - {s}")

# 主机总数
total_hosts = sum(int(table1.rows[i].cells[2].text) for i in range(1, len(table1.rows)-1))
print(f"\n主机总数: {total_hosts} 台")

# CPU总量
total_cpu = sum(int(table1.rows[i].cells[3].text) for i in range(1, len(table1.rows)-1))
print(f"CPU总量: {total_cpu} 核")

# 验证VPN匹配正确
if vpn_table:
    has_nonglin = any('农林' in vpn_table.rows[r].cells[1].text 
                      for r in range(1, min(10, len(vpn_table.rows))))
    print(f"\nVPN包含农林数据: {'是 ✓' if has_nonglin else '否 ✗'}")

print("\n" + "="*70)
print("✅✅✅ 北京市农林科学院标准月报验证完成！")
print("="*70)
print(f"\n输出文件: {files[0]}")
