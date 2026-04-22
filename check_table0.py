#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)

fp = os.path.join(output_dir, files[0])
doc = Document(fp)

table = doc.tables[0]
print(f"Table 0: {len(table.rows)} rows")

for i, row in enumerate(table.rows):
    print(f"\nRow {i}:")
    for j, cell in enumerate(row.cells):
        text = cell.text.strip()
        alignments = [p.alignment for p in cell.paragraphs]
        print(f"  Cell {j}: '{text[:20]}' alignments: {alignments}")
