#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""检查北京市卫生健康人力资源发展中心的表格数据"""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_175016-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 检查表格数据")
    print("="*70)
    
    # 检查CPU使用率表
    print("\n📊 CPU使用率汇总表 (表格1):")
    table1 = doc.tables[1]
    for row_idx, row in enumerate(table1.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行{row_idx}: {cells}")
    
    # 检查内存使用率表
    print("\n📊 内存使用率汇总表 (表格2):")
    table2 = doc.tables[2]
    for row_idx, row in enumerate(table2.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行{row_idx}: {cells}")
    
    # 检查存储使用率表
    print("\n📊 存储使用率汇总表 (表格3):")
    table3 = doc.tables[3]
    for row_idx, row in enumerate(table3.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  行{row_idx}: {cells}")
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
