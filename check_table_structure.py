#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

template_path = r'D:\月报自动化\月报模板\政务云服务运维月报-2025年11月-北京市农林科学院.docx'

try:
    doc = Document(template_path)
    print(f"文档段落数: {len(doc.paragraphs)}")
    print(f"文档表格数: {len(doc.tables)}")
    
    # 查找"本月技术支撑统计"相关的段落和表格
    print("\n查找'技术支撑'相关内容:")
    for i, para in enumerate(doc.paragraphs[:30]):
        if '技术支撑' in para.text or '工单' in para.text:
            print(f"\n段落 {i}: {para.text}")
    
    # 查看前几个表格的结构
    print("\n" + "="*60)
    print("查看前3个表格的结构:")
    for i, table in enumerate(doc.tables[:3]):
        print(f"\n--- 表格 {i} ---")
        print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
        
        # 打印表格内容
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"行{row_idx}: {cells}")
            if row_idx >= 5:  # 只显示前5行
                print("...")
                break
        
        # 检查是否是技术支撑统计表
        first_row_text = ' '.join([cell.text for cell in table.rows[0].cells])
        if '技术' in first_row_text or '工单' in first_row_text or '工作类型' in first_row_text:
            print(f">>> 这是技术支撑统计表 (表格{i})")

except Exception as e:
    print(f"错误: {e}")
    import traceback
    traceback.print_exc()
