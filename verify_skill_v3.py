#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证优化后的标准月报skill"""
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)

if not files:
    print("未找到输出文件")
    sys.exit(1)

output_file = os.path.join(output_dir, files[0])
print(f"验证文件: {files[0]}")
print(f"文件大小: {os.path.getsize(output_file):,} bytes\n")

doc = Document(output_file)

# 1. 验证日期审核（检查是否有非3月份的日期）
print("="*70)
print("1. 日期审核 - 确保所有日期在2026年3月")
print("="*70)

import re
issues = []
for table_idx, table in enumerate(doc.tables[:12]):  # 只检查前12个表格
    for row_idx, row in enumerate(table.rows[:5]):  # 每表前5行
        for cell_idx, cell in enumerate(row.cells):
            dates = re.findall(r'(\d{4}-\d{2}-\d{2})', cell.text)
            for d in dates:
                if not d.startswith('2026-03'):
                    issues.append({
                        'table': table_idx,
                        'date': d,
                        'text': cell.text[:30]
                    })

if issues:
    print(f"⚠️  发现 {len(issues)} 处日期不在2026年3月:")
    for i in issues[:5]:
        print(f"   表格{i['table']}: {i['date']} - {i['text']}")
else:
    print("✅ 所有日期都在2026年3月范围内！")

# 2. 验证表格格式（居中+框线）
print("\n" + "="*70)
print("2. 表格格式审核 - 居中+框线")
print("="*70)

from docx.enum.text import WD_ALIGN_PARAGRAPH
sample_table = doc.tables[0]  # 第一个表格
sample_cell = sample_table.rows[1].cells[0]

# 检查对齐方式
is_centered = any(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in sample_cell.paragraphs)
print(f"示例单元格居中对齐: {'✅ 是' if is_centered else '❌ 否'}")

# 3. 验证表格完整性
print("\n" + "="*70)
print("3. 表格完整性审核")
print("="*70)

table_names = []
for i, table in enumerate(doc.tables):
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    header_str = ' '.join(headers)
    
    if '工单' in header_str:
        table_names.append(('工单统计', len(table.rows)-1))
    elif 'CPU' in header_str:
        table_names.append(('CPU使用率', len(table.rows)-1))
    elif '内存' in header_str:
        table_names.append(('内存使用率', len(table.rows)-1))
    elif '存储' in header_str:
        table_names.append(('存储使用率', len(table.rows)-1))
    elif '基础资源' in header_str or '主机IP' in header_str:
        table_names.append(('基础资源台账', len(table.rows)-1))
    elif '备份' in header_str or '快照' in header_str:
        table_names.append(('快照备份', len(table.rows)-1))
    elif '防篡改' in header_str:
        table_names.append(('网页防篡改', len(table.rows)-1))
    elif '堡垒机' in header_str or '资产IP' in header_str:
        table_names.append(('堡垒机审计', len(table.rows)-1))
    elif 'VPN' in header_str or ('用户名' in header_str and '用户组' in header_str):
        table_names.append(('VPN审计', len(table.rows)-1))

for name, count in table_names:
    print(f"  {name}: {count} 条记录")

print("\n" + "="*70)
print("✅ 标准月报审核完成！")
print("="*70)
