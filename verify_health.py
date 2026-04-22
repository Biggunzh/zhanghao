#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_164433-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("✅ 文档验证")
    print(f"段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
    
    # 验证日期
    print("\n日期验证:")
    for i, para in enumerate(doc.paragraphs[:50]):
        if '2026年3月' in para.text and ('卫生' in para.text or '共有' in para.text):
            print(f"段落{i}: {para.text[:60]}...")
            break
    
    # 验证表1
    print("\n表1（技术支撑统计）验证:")
    table = doc.tables[0]
    print(f"表格: {len(table.rows)}行 x {len(table.columns)}列")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
except Exception as e:
    print(f"❌ 错误: {e}")
