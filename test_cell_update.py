#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def clear_and_set_cell(cell, value):
    """清空单元格现有数据并设置新值，加上居中+边框格式"""
    # 1. 设置单元格文本值
    cell.text = str(value)
    
    # 2. 居中对齐
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. 确保单元格有边框
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 设置单元格边框
    try:
        tcBorders = parse_xml(r'''
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tcBorders>
        ''')
        tcPr.append(tcBorders)
    except:
        pass

# 测试
doc = Document(r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03_v2_192203-北京市公安局勤务指挥部.docx')

# 找到快照备份表格
for i, table in enumerate(doc.tables):
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    header_str = ' '.join(headers)
    
    if '备份时间' in header_str:
        print(f"找到表格{i}: {headers}")
        
        # 测试修改第一行数据
        if len(table.rows) > 1:
            row = table.rows[1]
            backup_time_cell = row.cells[3]  # 备份时间列
            
            print(f"\n修改前: {backup_time_cell.text}")
            
            # 修改
            clear_and_set_cell(backup_time_cell, "2026/3/6 22:00")
            
            print(f"修改后: {backup_time_cell.text}")
        
        break

# 保存测试
doc.save(r'D:\月报自动化\test_output.docx')
print("\n已保存到 test_output.docx")
