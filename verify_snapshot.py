#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 找到最新文件
output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '农林' in f and f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
output_file = os.path.join(output_dir, files[0])

print(f"验证文件: {files[0]}")
print(f"文件大小: {os.path.getsize(output_file):,} bytes\n")

doc = Document(output_file)

print("="*70)
print("查找所有表格...")
print("="*70)

for i, table in enumerate(doc.tables):
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    header_str = ' '.join(headers)
    print(f"\n表格{i}: {len(table.columns)}列 x {len(table.rows)}行")
    print(f"  表头: {headers[:5]}{'...' if len(headers) > 5 else ''}")
    
    # 判断表格类型
    if '快照' in header_str or '备份' in header_str:
        print(f"  ⚠️ 可能是快照备份表！")
    elif '开始时间' in header_str and '资产IP' in header_str:
        print(f"  → 堡垒机审计表")
    elif '用户名' in header_str and '用户组' in header_str:
        print(f"  → VPN审计表")

print("\n" + "="*70)
print("查找快照备份服务报告表...")
print("="*70)

snapshot_table = None
for i, table in enumerate(doc.tables):
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    header_str = ' '.join(headers)
    
    # 查找快照备份表
    if ('云主机名称' in header_str or '主机名称' in header_str) and \
       ('备份时间' in header_str or '备份类型' in header_str):
        snapshot_table = table
        print(f"\n✅ 找到快照备份表（表格{i}）:")
        print(f"  列数: {len(table.columns)}, 行数: {len(table.rows)}")
        print(f"  表头: {headers}")
        break

if snapshot_table:
    print(f"\n前5行数据:")
    for row_idx in range(1, min(6, len(snapshot_table.rows))):
        cells = [cell.text.strip() for cell in snapshot_table.rows[row_idx].cells]
        print(f"  {cells}")
else:
    print("\n❌ 未找到快照备份服务报告表")
    print("\n可能原因：")
    print("  1. 模板中不存在此表格")
    print("  2. 表头字段不匹配")
    print("\n请检查模板文件中的附件2表格")
