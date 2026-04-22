#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_164021-司法局.docx'
doc = Document(output_file)

print("输出文件中所有包含'202'的段落:")
print("="*60)

for i, para in enumerate(doc.paragraphs):
    if '202' in para.text:
        text = para.text.strip()
        if text:
            print(f"\n段落{i}: {text[:100]}")

print("\n\n验证:")
print(f"段落23: {doc.paragraphs[23].text}")
print(f"段落44: {doc.paragraphs[44].text[:80]}...")
