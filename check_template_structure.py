#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

template = r'D:\月报自动化\月报模板\政务云服务运维月报-2025年11月-北京市卫生健康人力资源发展中心.docx'

doc = Document(template)

print("="*70)
print("检查模板表格结构")
print("="*70)

# 检查表格1 (CPU使用率汇总)
table1 = doc.tables[1]
print(f"\n表格1 (CPU使用率):")
print(f"  总行数: {len(table1.rows)}")
print(f"  总列数: {len(table1.columns)}")

# 检查是否有合并单元格
print("\n  检查合并单元格:")
for row_idx, row in enumerate(table1.rows[:5]):
    cells_info = []
    for col_idx, cell in enumerate(row.cells):
        # 检查是否是合并单元格的一部分
        tc = cell._tc
        grid_span = tc.grid_span if hasattr(tc, 'grid_span') else 1
        cells_info.append(f"{col_idx}(span={grid_span})")
    print(f"    行{row_idx}: {cells_info}")

# 打印前5行的实际内容
print("\n  前5行内容:")
for row_idx in range(min(5, len(table1.rows))):
    cells = [cell.text.strip() for cell in table1.rows[row_idx].cells]
    print(f"    行{row_idx}: {cells}")

# 同样检查存储表
table3 = doc.tables[3]
print(f"\n表格3 (存储使用率):")
print(f"  总行数: {len(table3.rows)}")
print(f"  总列数: {len(table3.columns)}")

print("\n  前5行内容:")
for row_idx in range(min(5, len(table3.rows))):
    cells = [cell.text.strip() for cell in table3.rows[row_idx].cells]
    print(f"    行{row_idx}: {cells}")
