#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证北京市卫生健康人力资源发展中心完整输出"""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_170013-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 北京市卫生健康人力资源发展中心 - 完整验证")
    print("="*70)
    
    # 1. 验证基础资源台账概况
    print("\n📋 1. 基础资源台账概况")
    print("-"*70)
    for para in doc.paragraphs[:50]:
        text = para.text.strip()
        if '共有' in text and '业务系统' in text:
            print(text)
            break
    
    # 2. 验证表1 - 本月技术支撑统计
    print("\n📋 2. 本月技术支撑统计（表1）")
    print("-"*70)
    table1 = doc.tables[0]
    for row in table1.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    # 3. 验证表2-4 - CPU/内存/存储使用率
    print("\n📋 3. 业务系统资源使用情况统计")
    print("-"*70)
    
    resource_tables = []
    for i, table in enumerate(doc.tables[1:5]):
        if len(table.rows) >= 3 and len(table.columns) >= 5:
            header_text = table.rows[0].cells[0].text.strip()
            if 'CPU' in header_text or '内存' in header_text or '存储' in header_text:
                resource_tables.append((i+1, table, header_text))
    
    for idx, table, header in resource_tables[:3]:
        print(f"\n  表格{idx} - {header[:10]}")
        # 显示表头
        header_row = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"  表头: {header_row}")
        # 显示数据行（限制前5行）
        for row_idx in range(1, min(6, len(table.rows))):
            cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
            if cells[0] and not cells[0].isdigit():
                continue
            print(f"    {cells}")
    
    # 4. 验证附件1开头
    print("\n📋 4. 附件1 - 政务云基础资源台账（前10行）")
    print("-"*70)
    for i, table in enumerate(doc.tables[4:8]):
        if len(table.columns) >= 6:
            # 可能是附件1
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if '序号' in header and '业务系统' in str(header):
                print(f"  表{i+5} 表头: {header[:7]}")
                for row_idx in range(1, min(11, len(table.rows))):
                    cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
                    print(f"    {cells[:5]}")  # 只显示前5列
                break
    
    print("\n" + "="*70)
    print("✅ 验证完成！")
    print("="*70)
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
