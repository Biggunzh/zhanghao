#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_164021-司法局.docx'

try:
    doc = Document(output_file)
    print("检查页眉/页脚中的日期...")
    
    for i, section in enumerate(doc.sections):
        print(f"\n--- 节 {i} ---")
        
        # 检查默认页眉
        if section.header:
            print("默认页眉:")
            for j, para in enumerate(section.header.paragraphs):
                if para.text.strip():
                    print(f"  段落{j}: {para.text}")
        
        # 检查首页页眉
        if section.first_page_header:
            print("首页页眉:")
            for j, para in enumerate(section.first_page_header.paragraphs):
                if para.text.strip():
                    print(f"  段落{j}: {para.text}")
        
        # 检查奇偶页眉
        if section.even_page_header:
            print("偶数页页眉:")
            for j, para in enumerate(section.even_page_header.paragraphs):
                if para.text.strip():
                    print(f"  段落{j}: {para.text}")
        
        # 检查页脚
        if section.footer:
            print("默认页脚:")
            for j, para in enumerate(section.footer.paragraphs):
                if para.text.strip():
                    print(f"  段落{j}: {para.text}")
    
    # 检查原始模板
    print("\n\n--- 检查原始模板 ---")
    template = r'D:\月报自动化\月报模板\政务云服务运维月报-2025年11月-司法局.docx'
    doc2 = Document(template)
    
    print("原始模板页眉:")
    for i, section in enumerate(doc2.sections[:1]):
        if section.header:
            for j, para in enumerate(section.header.paragraphs):
                if para.text.strip():
                    print(f"  {para.text}")
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
