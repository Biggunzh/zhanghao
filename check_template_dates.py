#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

template = r'D:\月报自动化\月报模板\政务云服务运维月报-2025年11月-司法局.docx'
doc = Document(template)

print("原始模板中所有包含'2025'的段落:")
print("="*60)

found = 0
for i, para in enumerate(doc.paragraphs):
    if '2025' in para.text or '11月' in para.text:
        text = para.text.strip()
        if text:
            print(f"\n段落{i}: {text}")
            found += 1

print(f"\n共找到 {found} 个包含2025的段落")

# 检查表格中的日期
print("\n\n表格中是否包含日期:")
for i, table in enumerate(doc.tables[:3]):
    for row in table.rows:
        for cell in row.cells:
            if '2025' in cell.text:
                print(f"表格{i}: {cell.text}")
