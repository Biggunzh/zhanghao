#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 找到最新司法局文件
output_dir = r'D:\月报自动化\输出月报'
files = [f for f in os.listdir(output_dir) if '司法局' in f and f.endswith('.docx') and not f.startswith('~$')]
files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
output_file = os.path.join(output_dir, files[0])

print(f"验证文件: {files[0]}\n")

doc = Document(output_file)

# 查找附件3 - 堡垒机审计记录表
print("="*70)
print("📋 附件3 - 堡垒机审计记录")
print("="*70)

fortress_table = None
for i, table in enumerate(doc.tables):
    if len(table.columns) >= 8 and len(table.rows) > 10:
        header = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
        if '开始时间' in header and '资产IP' in header and '协议' in header:
            fortress_table = table
            print(f"\n找到堡垒机审计记录表 (表格{i}):")
            print(f"  列数: {len(table.columns)}, 行数: {len(table.rows)}")
            break

if fortress_table:
    # 显示表头
    headers = [cell.text.strip() for cell in fortress_table.rows[0].cells]
    print(f"\n表头: {headers}")
    
    # 显示前5行数据
    print("\n前5行数据:")
    for row_idx in range(1, min(6, len(fortress_table.rows))):
        cells = [cell.text.strip() for cell in fortress_table.rows[row_idx].cells]
        print(f"  行{row_idx}: {cells}")
    
    # 显示最后3行
    print("\n最后3行数据:")
    for row_idx in range(max(1, len(fortress_table.rows)-3), len(fortress_table.rows)):
        cells = [cell.text.strip() for cell in fortress_table.rows[row_idx].cells]
        print(f"  行{row_idx}: {cells}")
    
    # 统计唯一的资产IP
    asset_ips = set()
    for row_idx in range(1, len(fortress_table.rows)):
        ip = fortress_table.rows[row_idx].cells[5].text.strip() if len(fortress_table.rows[row_idx].cells) > 5 else ''
        if ip:
            asset_ips.add(ip)
    
    print(f"\n✅ 统计:")
    print(f"  总记录数: {len(fortress_table.rows) - 1} 条")
    print(f"  涉及资产IP数: {len(asset_ips)} 个")
    print(f"  前10个IP: {list(asset_ips)[:10]}")
    
    # 验证格式（居中）
    print(f"\n✅ 格式检查:")
    sample_cell = fortress_table.rows[1].cells[5]  # 资产IP列
    alignment = sample_cell.paragraphs[0].alignment if sample_cell.paragraphs else None
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    is_centered = alignment == WD_ALIGN_PARAGRAPH.CENTER
    print(f"  数据是否居中: {'是 ✓' if is_centered else '否 ✗'}")
    
else:
    print("❌ 未找到堡垒机审计记录表")

print("\n" + "="*70)
