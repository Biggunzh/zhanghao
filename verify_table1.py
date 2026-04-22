#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_163222-北京市农林科学院.docx'

try:
    doc = Document(output_file)
    print("✅ 文档打开成功！")
    
    # 查看表1（技术支撑统计表）
    table = doc.tables[0]
    print(f"\n表1（技术支撑统计表）:")
    print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
    
    for row_idx, row in enumerate(table.rows):
        cells = row.cells
        cell_texts = []
        for cell in cells:
            text = cell.text.strip()
            # 检查对齐方式
            alignment = cell.paragraphs[0].alignment if cell.paragraphs else None
            cell_texts.append(f"'{text}'(对齐:{alignment})")
        print(f"行{row_idx}: {cell_texts}")
    
    # 验证数值
    data_cell = table.rows[1].cells[2]
    print(f"\n工单数量单元格内容: '{data_cell.text.strip()}'")
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
