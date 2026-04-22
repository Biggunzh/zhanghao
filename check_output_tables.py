#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_180233-北京市卫生健康人力资源发展中心.docx'

doc = Document(output_file)

print(f"文档中共有 {len(doc.tables)} 个表格\n")

for i, table in enumerate(doc.tables[:10]):
    # 获取表头
    headers = []
    for row in table.rows[:2]:
        row_text = ' '.join([c.text.strip() for c in row.cells])
        if row_text:
            headers.append(row_text[:50])
    
    print(f"表格{i}: {len(table.rows)}行 x {len(table.columns)}列")
    print(f"  表头: {headers}")
    print()
