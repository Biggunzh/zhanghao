#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
月报自动化处理脚本
功能：根据原始数据自动填充Word模板生成月报
"""

import os
import sys
import re
from datetime import datetime, timedelta
from collections import defaultdict

# 配置编码
sys.stdout.reconfigure(encoding='utf-8')

# ============ 配置区域 ============
BASE_DIR = r'D:\月报自动化'
RAW_DATA_DIR = os.path.join(BASE_DIR, '月报原始数据')
TEMPLATE_DIR = os.path.join(BASE_DIR, '月报模板')
OUTPUT_DIR = os.path.join(BASE_DIR, '输出月报')

# 原始数据文件
RESOURCE_FILE = os.path.join(RAW_DATA_DIR, '2026-03月报资源使用率详情列表.xls')
WORKORDER_FILE = os.path.join(RAW_DATA_DIR, '2026-03工单总量.xlsx')
FORTRESS_FILE = os.path.join(RAW_DATA_DIR, '2026-03-堡垒机.xlsx')
VPN_FILE = os.path.join(RAW_DATA_DIR, '2026-03vpn审计.xlsx')

# ============ Excel读取工具 ============
def read_excel_simple(file_path, max_rows=None):
    """简单读取Excel文件，返回列表的列表"""
    try:
        if file_path.endswith('.xls'):
            # 尝试使用xlrd
            try:
                import xlrd
                book = xlrd.open_workbook(file_path)
                sheet = book.sheet_by_index(0)
                data = []
                for i in range(sheet.nrows):
                    if max_rows and i >= max_rows:
                        break
                    row = [str(cell.value) if cell.value is not None else '' 
                           for cell in sheet.row(i)]
                    data.append(row)
                return data
            except ImportError:
                pass
        
        # 使用openpyxl
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        data = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if max_rows and i >= max_rows:
                break
            data.append([str(cell) if cell is not None else '' for cell in row])
        wb.close()
        return data
    except Exception as e:
        print(f"❌ 读取文件失败 {file_path}: {e}")
        return []

# ============ 数据提取函数 ============
def extract_business_systems_from_word(docx_path):
    """从Word模板中提取业务系统名称"""
    try:
        from docx import Document
        doc = Document(docx_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        
        # 查找业务系统名称（通常在"3 业务系统资源使用情况统计"部分）
        # 匹配类似 "微营销"、"长城网" 这样的系统名
        systems = set()
        
        # 从表格中提取
        for table in doc.tables[:3]:  # 只看前3个表格
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    # 检查是否是业务系统行
                    if cells[0] in ['1', '2', '3'] and cells[1]:
                        systems.add(cells[1])
        
        return sorted(list(systems))
    except Exception as e:
        print(f"❌ 提取业务系统失败: {e}")
        return []

def parse_resource_data(file_path):
    """解析资源使用率详情数据"""
    data = read_excel_simple(file_path)
    if not data:
        return {}
    
    # 找到表头行（通常包含"业务系统"、"主机名"等）
    header_row = None
    for i, row in enumerate(data[:15]):
        if '业务系统' in str(row) or '系统' in str(row):
            header_row = i
            break
    
    if header_row is None:
        print("⚠️ 未找到表头行，使用第一行作为表头")
        header_row = 0
    
    headers = data[header_row]
    print(f"📊 资源数据表头: {headers[:15]}")  # 打印前15列
    
    # 构建数据字典，按业务系统分组
    systems_data = defaultdict(lambda: {
        'hosts': [],
        'cpu_count': 0,
        'memory_gb': 0,
        'storage_gb': 0,
        'cpu_usage_sum': 0,
        'memory_usage_sum': 0,
        'storage_usage_sum': 0,
        'host_count': 0
    })
    
    for row in data[header_row+1:]:
        if not row or not row[0]:
            continue
        
        # 提取业务系统名称（通常在B列，索引1）
        system_name = str(row[1]).strip() if len(row) > 1 else ''
        if not system_name:
            continue
        
        # 提取主机信息（根据实际表头匹配列索引）
        # 列索引: 0=业务系统编号, 1=业务系统名称, 2=云主机ID, 3=云主机名称, 4=实例状态, 5=CPU, 6=内存(GB), 7=磁盘(总GB), 8=固定IP
        host_info = {
            'host_name': str(row[3]).strip() if len(row) > 3 else '',
            'ip': str(row[8]).strip() if len(row) > 8 else '',  # 固定IP在索引8
            'cpu': int(float(row[5])) if len(row) > 5 and row[5] else 0,
            'memory': int(float(row[6])) if len(row) > 6 and row[6] else 0,
            'storage': int(float(row[7])) if len(row) > 7 and row[7] else 0,
        }
        
        # 累计资源
        systems_data[system_name]['hosts'].append(host_info)
        systems_data[system_name]['cpu_count'] += host_info['cpu']
        systems_data[system_name]['memory_gb'] += host_info['memory']
        systems_data[system_name]['storage_gb'] += host_info['storage']
        systems_data[system_name]['host_count'] += 1
    
    return dict(systems_data)

def parse_workorder_data(file_path):
    """解析工单数据"""
    data = read_excel_simple(file_path)
    if not data:
        return {}
    
    # 按业务系统统计工单
    workorder_stats = defaultdict(lambda: {
        '需求处理': 0,
        '故障处理': 0,
        '咨询': 0,
        '其他': 0,
        'total': 0
    })
    
    # 找到表头
    header_row = None
    for i, row in enumerate(data[:10]):
        if any('业务系统' in str(cell) for cell in row):
            header_row = i
            break
    
    if header_row is None:
        header_row = 0
    
    for row in data[header_row+1:]:
        if len(row) < 2:
            continue
        
        system_name = str(row[0]).strip() if row[0] else ''
        workorder_type = str(row[1]).strip() if len(row) > 1 and row[1] else '其他'
        
        if not system_name:
            continue
        
        if '需求' in workorder_type:
            workorder_stats[system_name]['需求处理'] += 1
        elif '故障' in workorder_type:
            workorder_stats[system_name]['故障处理'] += 1
        elif '咨询' in workorder_type:
            workorder_stats[system_name]['咨询'] += 1
        else:
            workorder_stats[system_name]['其他'] += 1
        
        workorder_stats[system_name]['total'] += 1
    
    return dict(workorder_stats)

def parse_fortress_data(file_path):
    """解析堡垒机数据"""
    data = read_excel_simple(file_path)
    return data

def parse_vpn_data(file_path):
    """解析VPN审计数据"""
    data = read_excel_simple(file_path)
    return data

# ============ Word文档处理 ============
def replace_text_in_paragraph(paragraph, old_text, new_text):
    """替换段落中的文本"""
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

def update_table_cell(table, row_idx, col_idx, new_value):
    """更新表格单元格内容（保持格式）"""
    try:
        if row_idx < len(table.rows) and col_idx < len(table.columns):
            cell = table.rows[row_idx].cells[col_idx]
            # 保持原有格式，只更新文本
            cell.text = str(new_value)
    except Exception as e:
        print(f"⚠️ 更新表格单元格失败 [{row_idx},{col_idx}]: {e}")

def get_last_month():
    """获取上个月的年月"""
    today = datetime.now()
    if today.month == 1:
        last_month = 12
        year = today.year - 1
    else:
        last_month = today.month - 1
        year = today.year
    return year, last_month

def generate_monthly_report(template_path, output_path):
    """生成月报"""
    print(f"\n{'='*60}")
    print(f"📝 处理模板: {os.path.basename(template_path)}")
    print(f"{'='*60}")
    
    # 1. 读取原始数据
    print("\n📊 步骤1: 读取原始数据...")
    resource_data = parse_resource_data(RESOURCE_FILE)
    workorder_data = parse_workorder_data(WORKORDER_FILE)
    fortress_data = parse_fortress_data(FORTRESS_FILE)
    vpn_data = parse_vpn_data(VPN_FILE)
    
    print(f"   发现业务系统: {list(resource_data.keys())}")
    
    # 2. 提取客户名称
    customer_name = None
    template_name = os.path.basename(template_path)
    match = re.search(r'-([^-]+)\.docx$', template_name)
    if match:
        customer_name = match.group(1)
    print(f"   客户名称: {customer_name}")
    
    # 3. 打开Word模板
    print("\n📝 步骤2: 处理Word文档...")
    try:
        from docx import Document
        doc = Document(template_path)
    except Exception as e:
        print(f"❌ 打开Word文档失败: {e}")
        return False
    
    # 4. 获取上个月日期
    year, month = get_last_month()
    last_month_str = f"{year}年{month}月"
    
    # 5. 替换日期
    print("\n📅 步骤3: 替换日期...")
    for para in doc.paragraphs:
        # 匹配类似 "2025年11月" 的日期
        replace_text_in_paragraph(para, '2025年11月', last_month_str)
        replace_text_in_paragraph(para, '2025年12月', last_month_str)
    
    # 6. 计算总资源
    print("\n📊 步骤4: 计算资源数据...")
    total_hosts = 0
    total_cpu = 0
    total_memory = 0
    total_storage = 0
    
    for system_name, data in resource_data.items():
        total_hosts += data['host_count']
        total_cpu += data['cpu_count']
        total_memory += data['memory_gb']
        total_storage += data['storage_gb']
    
    print(f"   总主机数: {total_hosts}")
    print(f"   总CPU: {total_cpu}核")
    print(f"   总内存: {total_memory}GB")
    print(f"   总存储: {total_storage}GB")
    
    # 7. 替换基础资源概况中的数据
    print("\n📝 步骤5: 更新基础资源台账概况...")
    for para in doc.paragraphs:
        text = para.text
        # 替换总资源数（需要找到包含"共有""台主机"的段落）
        if '共有' in text and '台主机' in text:
            # 使用正则替换数字
            # 例如: "共有2个业务系统...共使用121台主机，899颗CPU，2899GB内存，75646GB存储"
            # 替换为实际计算的值
            new_text = text
            new_text = re.sub(r'共使用\d+台主机', f'共使用{total_hosts}台主机', new_text)
            new_text = re.sub(r'\d+颗CPU', f'{total_cpu}颗CPU', new_text)
            new_text = re.sub(r'\d+GB内存', f'{total_memory}GB内存', new_text)
            new_text = re.sub(r'\d+GB存储', f'{total_storage}GB存储', new_text)
            
            if new_text != text:
                for run in para.runs:
                    run.text = new_text
                print(f"   已更新资源概况段落")
    
    # 8. 更新表格数据
    print("\n📋 步骤6: 更新表格数据（简化处理）...")
    system_names = list(resource_data.keys())
    
    # 更新前几个关键表格
    table_count = 0
    for table in doc.tables[:6]:  # 处理前6个表格
        table_count += 1
        print(f"   处理表格 {table_count} (行数: {len(table.rows)}, 列数: {len(table.columns)})")
        
        # 如果表格有业务系统列，更新对应数据
        for i, row in enumerate(table.rows):
            if i == 0:
                continue  # 跳过表头
            
            cells = row.cells
            if len(cells) < 2:
                continue
            
            cell_text = cells[1].text.strip()  # 第二列通常是业务系统名称
            
            # 检查是否是已知的业务系统
            if cell_text in resource_data:
                data = resource_data[cell_text]
                # 更新后续列（CPU、内存、存储、使用率等）
                # 注意：这里需要根据实际表格结构调整列索引
                try:
                    if len(cells) > 2:
                        cells[2].text = str(data['host_count'])
                    if len(cells) > 3:
                        cells[3].text = str(data['cpu_count'])
                    if len(cells) > 4:
                        cells[4].text = '0.00'  # 使用率需要计算
                except:
                    pass
    
    # 9. 保存文档
    print(f"\n💾 步骤7: 保存文档到 {output_path}...")
    try:
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"   ✅ 文档已保存: {output_path}")
        return True
    except Exception as e:
        print(f"❌ 保存文档失败: {e}")
        return False

def process_all_templates():
    """处理所有模板"""
    print("\n" + "="*60)
    print("🚀 月报自动化处理系统")
    print("="*60)
    
    # 确保输出目录存在
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 获取所有模板文件
    template_files = [f for f in os.listdir(TEMPLATE_DIR) 
                      if f.endswith('.docx') and not f.startswith('~$')]
    
    print(f"\n📁 发现 {len(template_files)} 个模板文件")
    
    # 只测试第一个模板
    test_template = '政务云服务运维月报-2025年11月-北京市农林科学院.docx'
    if test_template in template_files:
        print(f"\n🧪 测试处理: {test_template}")
        template_path = os.path.join(TEMPLATE_DIR, test_template)
        output_path = os.path.join(OUTPUT_DIR, test_template.replace('2025年11月', '2026年03月'))
        generate_monthly_report(template_path, output_path)
    
    print("\n" + "="*60)
    print("✅ 处理完成")
    print("="*60)

if __name__ == '__main__':
    process_all_templates()
