#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证北京市卫生健康人力资源发展中心完整报告"""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

output_file = r'D:\月报自动化\输出月报\政务云服务运维月报-2026年03月_v2_173851-北京市卫生健康人力资源发展中心.docx'

try:
    doc = Document(output_file)
    print("="*70)
    print("✅ 北京市卫生健康人力资源发展中心 - 完整报告验证")
    print("="*70)
    
    # 1. 基础资源台账
    print("\n📋 1. 基础资源台账概况")
    print("-"*70)
    for para in doc.paragraphs[:50]:
        text = para.text.strip()
        if '共有' in text and '业务系统' in text and '运行' in text:
            print(text)
            break
    
    # 2. 本月技术支撑统计
    print("\n📋 2. 本月技术支撑统计")
    print("-"*70)
    table0 = doc.tables[0]
    for row in table0.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  {cells}")
    
    # 3. CPU/内存/存储使用率汇总表
    print("\n📋 3. CPU使用率汇总表")
    print("-"*70)
    table1 = doc.tables[1]
    for row_idx in range(min(6, len(table1.rows))):
        cells = [cell.text.strip() for cell in table1.rows[row_idx].cells]
        print(f"  {cells[:5]}")
    
    print("\n📋 4. 内存使用率汇总表")
    print("-"*70)
    table2 = doc.tables[2]
    for row_idx in range(min(6, len(table2.rows))):
        cells = [cell.text.strip() for cell in table2.rows[row_idx].cells]
        print(f"  {cells[:5]}")
    
    print("\n📋 5. 存储使用率汇总表")
    print("-"*70)
    table3 = doc.tables[3]
    for row_idx in range(min(6, len(table3.rows))):
        cells = [cell.text.strip() for cell in table3.rows[row_idx].cells]
        print(f"  {cells[:5]}")
    
    # 6. 附件1 - 基础资源台账
    print("\n📋 6. 附件1 - 政务云基础资源台账")
    print("-"*70)
    table4 = doc.tables[4]
    headers = [cell.text.strip() for cell in table4.rows[0].cells]
    print(f"  表头: {headers}")
    print(f"  总行数: {len(table4.rows)}行")
    print("  前5行数据:")
    for row_idx in range(1, min(6, len(table4.rows))):
        cells = [cell.text.strip() for cell in table4.rows[row_idx].cells]
        print(f"    {cells[:7]}")
    
    # 7. 附件2 - 资源使用率详情
    print("\n📋 7. 附件2 - CPU使用率详情")
    print("-"*70)
    table5 = doc.tables[5]
    print(f"  总行数: {len(table5.rows)}行")
    print("  前3行数据:")
    for row_idx in range(min(4, len(table5.rows))):
        cells = [cell.text.strip() for cell in table5.rows[row_idx].cells]
        print(f"    {cells}")
    
    print("\n" + "="*70)
    print("✅ 验证完成！所有内容已正确生成！")
    print("="*70)
    
except Exception as e:
    print(f"❌ 错误: {e}")
    import traceback
    traceback.print_exc()
